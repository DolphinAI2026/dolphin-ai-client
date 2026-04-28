"""需求分析路由 — 多轮对话式需求澄清 + 结构化功能设计文档生成"""
from __future__ import annotations

import base64
import json
import logging
import os
import re
import tempfile
from typing import Annotated, Optional, Any

from fastapi import APIRouter, Body, Depends, File, Form, HTTPException, UploadFile
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sse_starlette.sse import EventSourceResponse

import httpx
from app.database import get_db
from app.deps import get_auth_context, AuthContext
from app.llm_client import LLMClient
from app.models import Conversation, Message
from app.services.design_doc_preflight import validate_design_doc_preflight

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/requirements", tags=["需求分析"])

# ─── System Prompts ───────────────────────────────────────────────────────────

REQUIREMENTS_CHAT_PROMPT = """你是一位经验丰富的产品分析师，负责通过对话帮用户把业务需求梳理清楚，最终生成结构化的功能设计文档。

## 需求收集框架

按以下 6 个维度依次收集，每个维度收集完后归纳确认，再进入下一个：

**① 项目目标**
- 这个系统要解决什么业务问题？
- 主要使用场景和预期效果是什么？

**② 角色**
- 谁会使用这个系统？各有什么不同的职责？
- 【注意】平台已内置组织架构，"全部员工/直属上级"不需要单独定义为角色——只需识别有独立权限配置差异的业务角色（如：HR管理员、财务专员、审批专员、系统管理员等）。当用户说"员工都能提交"时，记录为"全员可操作，数据范围=本人"，不创建员工角色。

**③ 枚举值**
- 系统中哪些字段是下拉选项/状态值？（如：假期类型、申请状态、费用类别等）
- 每个枚举有哪些选项值？
- 【注意】审批状态（待审批/已通过/已拒绝）由平台流程引擎内置，不需要用户定义。

**④ 业务对象**
- 系统要管理哪些核心业务单据或实体？（如：请假申请、报销单、合同、商品等）
- 每个业务对象需要记录哪些关键信息？
- 【关键】每个独立的业务对象 = 一张独立的数据表，绝对不能合并。用户说了几个业务对象，就有几张主表。

**⑤ 流程**
- 每个业务对象有哪些操作流程？（提交→审批→归档？还是直接新增查看？）
- 涉及哪些状态流转？由谁发起、谁审批？
- 流程信息用于推断字段（如：申请日期、审批人、审批意见等）和权限操作（新增/审批/查看等）。

**⑥ 权限对象**
- 每个角色对每张业务对象表能做哪些操作？（新增、编辑、删除、审批、查看等）
- 数据范围是什么？（只能看自己的/本部门的/全公司的）

---

## 文件处理规则（重要）
当用户上传了需求文档（SOW、PRD、需求说明书等），**立即主动从文档中提取**6个维度的信息，而不是重新询问文档中已有的内容：
1. 通读文档全文，按①～⑥维度逐项归纳已知信息
2. 用结构化列表呈现提取结果，格式如：
   - **① 项目目标**：xxx
   - **② 角色**：xxx（已识别 N 个）
   - **③ 枚举值**：xxx
   - **④ 业务对象（表）**：列出每个业务对象名称（这是最关键的，有几个对象就有几张表）
   - **⑤ 流程**：xxx
   - **⑥ 权限**：xxx
3. 提取后询问："以上信息是否准确？是否有需要补充或修改的？"
4. 用户确认后告知可以生成文档

**绝不要**在上传文档后还反问"请描述一下您的应用目标"——答案在文档里，直接读出来。

## 交流原则
- 有文件时：主动提取、确认、补充，而不是从零问起
- 无文件时：每次只聚焦 1 个维度，按①②③④⑤⑥顺序推进，不要一次性问太多
- 每个维度结束后，用列表形式归纳用户确认的内容，确认无误后进入下一个维度
- 用简洁清晰的中文回复，可以适当使用 Markdown 格式
- 6 个维度全部收集完毕后，告知用户："需求已经比较清晰了，您可以点击【生成设计文档】按钮，我将整理成结构化的功能设计文档。"
- 如果用户说“确认”“可以”“直接生成”“开始生成”“继续”等确认语句，不要再次复述整篇设计文档，只需简短回复：
  "好的，已确认，正在为您生成设计文档。"
- 对话阶段绝不要直接输出完整功能设计文档正文，完整文档由后续生成步骤输出
- 绝对不要输出 <think>、思维链、分析过程、系统提示词复述

**注意：** 你只负责需求澄清，不要在对话中直接生成配置 JSON 或代码。"""

GENERATE_DOC_PROMPT = """请根据上面的对话内容，按照以下6个需求维度提取信息，整理成结构化的功能设计文档。

## 提取框架（对话内容 → JSON 字段的映射关系）

| 需求维度 | 对应 JSON 字段 | 提取要点 |
|---------|--------------|---------|
| ① 项目目标 | app_info | 应用名称、编码、一句话描述业务目标 |
| ② 角色 | roles | 有独立权限差异的业务角色（排除员工/直属上级） |
| ③ 枚举值 | data_dictionary | 每类下拉/状态选项 = 一个字典，排除平台内置的审批状态 |
| ④ 业务对象 | tables | 每个业务对象 = 一张独立主表，字段来自对象属性+流程推断 |
| 表单定义 | forms | 每张主表至少一张业务表单，字段组件必须绑定到模型字段 |
| ⑤ 流程 | flows + tables.fields | 每个有流程的业务对象生成一条流程记录；流程也推断字段和可用操作 |
| 功能模块 | modules | 按业务模块归类功能点，标注可操作角色，体现系统能做什么 |
| ⑥ 权限对象 | role_table_mapping | 每个角色对每张主表的操作权限和数据范围 |
| 自开发定义 | 暂不输出 | 当前阶段只沉淀低代码配置 MD；复杂自开发后续从 Vibe Coding/IDE 入口处理 |

---

必须输出严格的 JSON，不包含任何 Markdown 代码块标记、解释性文字或注释。直接输出 JSON 对象：

{
  "app_info": {
    "code": "应用编码（英文大写缩写，如 OMS、CRM、ERP，若未明确则根据业务域推断）",
    "name": "应用名称（中文）",
    "description": "应用简要描述（1-2句话，体现项目目标）"
  },
  "roles": [
    {
      "role_code": "角色编码（英文小写，如 hr_manager、dept_manager、finance_specialist）",
      "role_name": "角色名称（中文）",
      "description": "角色职责描述"
    }
  ],
  "data_dictionary": [
    {
      "dict_code": "字典编码（英文下划线，如 leave_type、expense_category）",
      "dict_name": "字典名称（中文，如 假期类型、费用类别）",
      "items": [
        {
          "item_code": "枚举值编码（英文大写，如 ANNUAL、SICK）",
          "item_name": "枚举值名称（中文，如 年假、病假）"
        }
      ]
    }
  ],
  "tables": [
    {
      "table_code": "表名（英文下划线，以 t_ 开头，如 t_leave_application）",
      "table_name": "表中文名称（如 请假申请）",
      "table_type": "主表 或 子表（二选一）",
      "parent_table": "父表 table_code（子表填写，主表填空字符串）",
      "description": "表用途描述",
      "fields": [
        {
          "field_code": "字段名（英文下划线）",
          "field_name": "字段中文名",
          "data_type": "VARCHAR/BIGINT/INT/DECIMAL/DATE/DATETIME/TINYINT/TEXT/BOOLEAN",
          "length": "长度（无则填空字符串）",
          "is_pk": false,
          "is_fk": false,
          "nullable": true,
          "default_value": "",
          "description": "字段描述"
        }
      ]
    }
  ],
  "forms": [
    {
      "form_code": "表单编码，默认与主表 table_code 保持一致",
      "form_name": "表单名称（中文，如 请假申请、巡检记录）",
      "model_code": "绑定主表 table_code",
      "all_model_codes": ["绑定的主表和子表 table_code"],
      "components": [
        {
          "field_code": "绑定字段编码",
          "field_name": "字段中文名",
          "component_type": "FORM_TEXT_INPUT/FORM_TEXTAREA/FORM_SELECT/FORM_DATE_PICKER/FORM_NUMBER_INPUT/FORM_UPLOAD/FORM_USER_SELECT/FORM_DEPT_SELECT",
          "section_type": "main 或 sub",
          "model_code": "字段所属模型 table_code",
          "required": true,
          "hidden": false,
          "readonly": false,
          "show_in_list": true,
          "searchable": false,
          "dict_code": "如字段绑定字典则填写字典编码，否则空字符串",
          "description": "组件说明"
        }
      ]
    }
  ],
  "modules": [
    {
      "module_name": "模块名称（中文，如 请假管理、加班管理、年假管理）",
      "module_code": "模块编码（英文下划线，如 leave_mgmt）",
      "description": "模块简要说明，1句话",
      "features": [
        {
          "name": "功能点名称（如 提交请假申请、查看申请状态、审批请假）",
          "description": "功能描述，说明做什么、达到什么效果",
          "roles": ["可操作此功能的角色名（用中文，如 全部员工、部门经理、HR专员）"]
        }
      ]
    }
  ],
  "flows": [
    {
      "flow_name": "流程名称（如 请假审批流程、加班申请流程）",
      "flow_code": "流程编码（英文下划线，如 leave_approval_flow）",
      "description": "流程整体说明",
      "steps": [
        {
          "step": 1,
          "action": "步骤动作（如 员工填写并提交申请）",
          "role": "执行该步骤的角色（用中文，如 全部员工、直属上级、HR专员）",
          "status": "该步骤后的状态（如 待审批、审批通过、已归档）"
        }
      ]
    }
  ],
  "role_table_mapping": [
    {
      "table_code": "主表表名",
      "table_name": "主表中文名",
      "permissions": [
        {
          "role_code": "角色编码",
          "role_name": "角色名称",
          "operations": ["从固定列表选择：暂存、新增、导入、复制新建、批量删除、批量同意、批量拒绝、查看、编辑、删除、查看审批历史、打印、日志、评论、导出"],
          "data_scope": "none/self/dept/all/custom（none=无权限、self=仅本人、dept=本部门、all=全公司、custom=自定义）"
        }
      ]
    }
  ],
  "custom_development": []
}

---

## 生成规则

### 【角色 roles】
- 只列有独立权限配置差异的业务角色
- 【严格禁止】泛化用户角色：员工、普通员工、全部员工、employee、staff、user、all_user（由平台组织架构管理）
- 【严格禁止】层级关系角色：直属上级、上级领导、leader、supervisor（由平台审批流程处理）
- 【严格禁止】把"审批人/approver"单独列为角色——审批是流程动作，应识别出实际承担审批职责的具体岗位（如 dept_manager 部门经理、hr_specialist HR专员），用该岗位角色命名；示例 role_code 中禁止出现 approver
- 这几类的权限通过 data_scope 和流程配置体现，不单独创建角色

### 【枚举值 data_dictionary】
- 每类下拉选项 = 一个字典，items 列出所有枚举值
- 【严格禁止】审批相关字典（平台内置）：approval_status、audit_status、process_status 等
- 从对话中的③枚举值维度提取，同时检查④业务对象的字段中是否还有遗漏的枚举

### 【数据表 tables — 最关键，必须完整提取】

**★ 第一步：扫描文档中所有业务实体，按以下五类逐一识别**

| 类别 | 说明 | 常见示例 |
|-----|------|---------|
| ① 申请/单据类 | 员工需要提交并经过审批的各种申请，每种独立建表 | 请假申请、加班申请、补卡申请、年假延期申请 |
| ② 配置/策略类 | 管理员可维护的规则、政策、参数，每类独立建表 | 假期政策配置、节假日设置、假期规则配置 |
| ③ 余额/台账类 | 追踪员工某类指标的额度或余量，每类独立建表 | 年假余额、加班余额、调休余额 |
| ④ 记录/归档类 | 业务完成后的结果记录，区别于"申请过程" | 考勤打卡记录、假期消耗记录 |
| ⑤ 扩展档案类 | 员工属性扩展或多对多关系表 | 员工假期档案、审批委托记录 |

**★ 第二步：每个识别到的业务实体 = 一张独立主表，绝对禁止合并**
- 禁止把请假/加班/补卡/调休合并成"一张申请表+类型字段"
- 禁止把政策/规则合并成"一张配置表+配置类型字段"
- 如果文档提到"X申请"和"Y申请"，两者各建一张表

**★ 第三步：每张表字段来源（综合三路）**
- 来自业务对象属性：该对象本身需要记录的信息（名称、金额、日期范围、备注等）
- 来自枚举字段：凡用到 data_dictionary 中字典的字段（如 leave_type → VARCHAR，存储枚举 code）
- 来自流程推断：有审批流程的表需要推断：申请日期(apply_date)、申请原因(reason)、附件说明(attachment_desc)等
- 来自关联关系：子表需要父表外键字段（is_fk=true）
- 【禁止】通用字段（平台自动管理）：id、create_time、update_time、create_by、update_by、deleted、approval_status

**★ 字段数量要求：每张表至少 6 个业务字段**

**★ 如果对话中包含 [上传文件：...] 的文件内容，必须扫描文件全文，将所有提及的业务实体（不论对话中是否明确确认）全部识别为表；文件内容中的表格、功能列表、业务流程描述都是提取依据**

### 【表单 forms】
- 每张主表必须生成至少 1 张表单，form_code 默认等于主表 table_code
- components 必须来自 tables.fields，不能出现没有模型字段支撑的组件
- 字段组件类型要按字段语义选择：文本用 FORM_TEXT_INPUT/FORM_TEXTAREA，日期用 FORM_DATE_PICKER，金额/数量用 FORM_NUMBER_INPUT，字典字段用 FORM_SELECT，附件用 FORM_UPLOAD
- 每个表单至少包含主表 6 个核心业务字段；常用查询字段 show_in_list/searchable 要合理标记
- 子表字段通过 section_type=sub、model_code=子表 table_code 表达

### 【权限矩阵 role_table_mapping】
- 只含主表，不含子表
- 每张主表的 permissions 数组，第一条必须是"全部员工"条目：`{"role_code":"all_employee","role_name":"全部员工","operations":[...],"data_scope":"self"}`，表示所有员工对该表的默认权限（通常是新增+暂存+查看自己的数据）
- 之后再列各业务角色的权限条目
- operations 根据⑤流程和⑥权限对象推断：有审批流程的表包含"批量同意/批量拒绝/查看审批历史"
- data_scope 推断：全部员工→self，经办人/申请人→self，部门管理者→dept，HR/管理员/全局角色→all

### 【功能模块 modules】
- 按业务域划分模块，每个主表通常对应1个模块（如"请假申请表" → "请假管理"模块）
- 每个模块的 features 列出该模块下所有功能点（增/删/改/查/审批/统计/导出等）
- roles 字段用中文角色名，全员可用时填"全部员工"，需要特定角色时填对应角色名
- 功能模块数量 = 主表数量（每张主表 ≥ 3 个功能点），另可有跨表的统计/报表模块

### 【业务流程 flows】
- 凡有审批/多步骤操作的业务对象，必须生成对应流程
- 每条流程 steps 按实际操作顺序排列，step 从 1 开始递增
- role 用中文描述（全部员工/直属上级/部门经理/HR专员等），status 描述该步完成后系统状态
- 纯 CRUD 无审批的表不需要生成流程

### 【自开发定义 custom_development】
- 当前阶段不生成自开发配置，custom_development 必须输出空数组 []
- 复杂前端组件、外部接口、Hook、插件、报表看板等内容只在需求摘要中记录为后续事项，不进入本次低代码配置 JSON
- 本次 JSON 只用于生成标准 MD 与低代码应用配置，避免把自开发内容误送到平台创建接口

### 【其他规则】
- is_pk、is_fk、nullable 必须是布尔值 true 或 false（不能是字符串）
- role_table_mapping 的 operations 只能从固定列表中选择，不能自创操作名"""

# ─── Tenant LLM Helpers ──────────────────────────────────────────────────────

async def _get_tenant_llm_config(db: AsyncSession, tenant_id: int) -> dict | None:
    """查询租户默认 LLM 配置，返回解密后的 dict"""
    from app.harness.llm_resolver import resolve_llm_config
    resolved = await resolve_llm_config(db, tenant_id, purpose="builder")
    if not resolved:
        return None
    return {"api_key": resolved.api_key, "base_url": resolved.base_url, "model": resolved.model, "max_tokens": resolved.max_tokens}


async def _get_conversation_llm_config(db: AsyncSession, conversation: Conversation) -> dict | None:
    from app.harness.llm_resolver import resolve_llm_config
    resolved = await resolve_llm_config(db, conversation.tenant_id, purpose="builder", selected_config_id=conversation.selected_llm_config_id)
    if not resolved:
        return None
    return {"api_key": resolved.api_key, "base_url": resolved.base_url, "model": resolved.model, "max_tokens": resolved.max_tokens}


async def _stream_with_config(cfg: dict | None, messages: list, max_retries: int = 2):
    """使用租户配置流式调用；Anthropic 兼容代理和 OpenAI 兼容接口都支持。"""
    if cfg is None:
        raise ValueError("未配置可用的 LLM 模型，请在环境管理 → 模型配置中添加并启用模型")

    llm = LLMClient(api_key=cfg["api_key"], base_url=cfg["base_url"], model=cfg["model"])

    last_err = None
    for attempt in range(max_retries + 1):
        if attempt > 0:
            import asyncio
            await asyncio.sleep(1)
            logger.info("LLM stream retry %d/%d for %s", attempt, max_retries, cfg["base_url"])
        try:
            async for chunk in llm.chat_completion_stream(messages, max_tokens=cfg.get("max_tokens", 8192)):
                yield chunk
            return  # 成功，不再重试
        except (httpx.ConnectError, httpx.ConnectTimeout, httpx.ReadTimeout) as e:
            last_err = e
            logger.warning("LLM stream attempt %d failed: %s", attempt + 1, e)
            continue
        except Exception:
            raise  # 非网络错误直接抛出

    # 所有重试都失败
    if last_err:
        raise last_err


async def _complete_with_config(
    cfg: dict | None,
    messages: list[dict[str, Any]],
    *,
    max_tokens: int = 4096,
    temperature: float = 0.0,
) -> str:
    """非流式调用，用于兜底修复 JSON。"""
    if cfg is None:
        raise ValueError("未配置可用的 LLM 模型")

    llm = LLMClient(api_key=cfg["api_key"], base_url=cfg["base_url"], model=cfg["model"])
    data = await llm.chat_completion(
        messages,
        max_tokens=max_tokens,
        timeout=120.0,
        temperature=temperature,
    )
    return ((data.get("choices") or [{}])[0].get("message") or {}).get("content", "")


async def _repair_doc_json(cfg: dict | None, raw_text: str) -> dict:
    """把模型输出的非标准/不完整文本修复成合法文档 JSON。"""
    repair_messages = [
        {
            "role": "system",
            "content": (
                "你是严格的 JSON 修复器。"
                "请把用户提供的文本修复为一个完整、合法的 JSON 对象。"
                "只输出 JSON，不要输出任何解释、Markdown、代码块、<think>。"
                "JSON 必须包含并仅按以下核心结构组织："
                "app_info(object), roles(array), data_dictionary(array), tables(array), "
                "modules(array), flows(array), role_table_mapping(array), forms(array), custom_development(array)。"
                "每张主表至少保留 6 个业务字段；每张主表都要有表单、功能模块和权限矩阵。"
                "当前配置阶段不生成自开发项，custom_development 必须是空数组 []。"
            ),
        },
        {
            "role": "user",
            "content": f"请修复以下内容并输出合法 JSON：\n\n{raw_text[-20000:]}",
        },
    ]
    repaired_text = await _complete_with_config(cfg, repair_messages, max_tokens=6000, temperature=0.0)
    return extract_json(repaired_text)


async def _regenerate_doc_json(cfg: dict | None, base_messages: list[dict[str, Any]]) -> dict:
    """当提取/修复都失败时，强约束再生成一次完整 JSON。"""
    regen_messages = list(base_messages)
    regen_messages.append(
        {
            "role": "user",
            "content": (
                "请忽略之前的输出形式，重新生成一次。"
                "严格只输出一个完整合法 JSON 对象，不能有 <think>、解释、注释、Markdown 代码块。"
                "必须包含 app_info、roles、data_dictionary、tables、modules、flows、"
                "role_table_mapping、forms、custom_development。"
                "每个业务对象独立成表，每张主表至少 6 个业务字段，并生成对应表单、流程和权限。"
                "custom_development 必须是空数组 []。如果无法确定字段值，给出合理默认值，但必须保证 SPEC 结构完整。"
            ),
        }
    )
    regen_text = await _complete_with_config(cfg, regen_messages, max_tokens=8000, temperature=0.0)
    return extract_json(regen_text)


# ─── Helpers ─────────────────────────────────────────────────────────────────

def extract_json(text: str) -> dict:
    """从 AI 响应中提取 JSON 对象"""
    # 去除 markdown 代码块
    text = re.sub(r'^```(?:json)?\s*', '', text.strip(), flags=re.MULTILINE)
    text = re.sub(r'\s*```$', '', text.strip(), flags=re.MULTILINE)
    text = text.strip()

    start = text.find('{')
    if start == -1:
        raise ValueError('响应中未找到 JSON 对象')

    depth = 0
    end = -1
    for i, ch in enumerate(text[start:], start):
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                end = i
                break

    if end == -1:
        raise ValueError('JSON 对象不完整')

    return json.loads(text[start:end + 1])


DOC_REQUIRED_LIST_KEYS = (
    "roles",
    "data_dictionary",
    "tables",
    "modules",
    "flows",
    "role_table_mapping",
    "forms",
)

SYSTEM_FIELD_CODES = {
    "id",
    "pk",
    "create_time",
    "created_at",
    "update_time",
    "updated_at",
    "create_by",
    "created_by",
    "update_by",
    "updated_by",
    "deleted",
    "deleted_at",
    "tenant_id",
    "org_id",
    "approval_status",
    "audit_status",
    "process_status",
}

FORBIDDEN_ROLE_CODES = {
    "employee",
    "staff",
    "user",
    "all_user",
    "all_employee",
    "approver",
    "leader",
    "supervisor",
}

FORBIDDEN_ROLE_NAME_KEYWORDS = ("员工", "全部员工", "普通用户", "审批人", "直属上级", "上级领导")


def _list_or_empty(value: Any) -> list:
    return value if isinstance(value, list) else []


def _text_from_messages(messages: list[dict[str, Any]] | None) -> str:
    if not messages:
        return ""
    chunks: list[str] = []
    for msg in messages:
        content = msg.get("content")
        if isinstance(content, str):
            chunks.append(content)
    return "\n".join(chunks)


def _infer_doc_app_name(messages: list[dict[str, Any]] | None = None) -> str:
    source = _text_from_messages(messages)
    match = re.search(r"([^\n，。；：]{2,30}(系统|平台|管理|应用))", source)
    if match:
        return match.group(1).strip()
    return "业务管理系统"


def _infer_doc_app_code(app_name: str) -> str:
    hints = [
        ("巡检", "insp_mgmt"),
        ("会议", "meeting_req"),
        ("项目", "project_mgmt"),
        ("客户", "crm"),
        ("CRM", "crm"),
        ("财务", "finance_mgmt"),
        ("停车", "parking_mgmt"),
        ("报销", "expense_mgmt"),
        ("请假", "leave_mgmt"),
    ]
    for keyword, code in hints:
        if keyword in app_name:
            return code
    ascii_code = re.sub(r"[^a-zA-Z0-9_]", "_", app_name).strip("_").lower()
    ascii_code = re.sub(r"_+", "_", ascii_code)[:24].strip("_")
    return ascii_code or "app_builder"


def _normalize_doc_app_info(raw: Any, messages: list[dict[str, Any]] | None = None) -> dict[str, str]:
    raw = raw if isinstance(raw, dict) else {}
    app_name = str(raw.get("name") or raw.get("app_name") or "").strip() or _infer_doc_app_name(messages)
    app_code = str(raw.get("code") or raw.get("app_code") or "").strip() or _infer_doc_app_code(app_name)
    description = str(raw.get("description") or "").strip() or f"{app_name}用于承载核心业务数据、流程、权限和后续扩展能力。"
    return {"code": app_code, "name": app_name, "description": description}


def _normalize_doc_data_type(raw_type: Any, field_name: str = "") -> str:
    value = str(raw_type or "").strip()
    if value:
        upper = value.upper()
        if upper in {"VARCHAR", "BIGINT", "INT", "DECIMAL", "DATE", "DATETIME", "TINYINT", "TEXT", "BOOLEAN"}:
            return upper
        mapping = {
            "单据号": "VARCHAR",
            "单行输入": "VARCHAR",
            "多行输入": "TEXT",
            "富文本": "TEXT",
            "下拉单选": "VARCHAR",
            "下拉多选": "VARCHAR",
            "单选框": "VARCHAR",
            "复选框": "VARCHAR",
            "日期": "DATE",
            "日期时间": "DATETIME",
            "时间": "DATETIME",
            "金额": "DECIMAL",
            "数字": "INT",
            "附件上传": "TEXT",
            "人员选择": "VARCHAR",
            "部门选择": "VARCHAR",
            "开关": "BOOLEAN",
        }
        if value in mapping:
            return mapping[value]
    if any(k in field_name for k in ("日期", "时间")):
        return "DATETIME"
    if any(k in field_name for k in ("金额", "数量", "次数", "评分", "百分比")):
        return "DECIMAL"
    if any(k in field_name for k in ("说明", "描述", "备注", "原因", "内容")):
        return "TEXT"
    return "VARCHAR"


def _field_template(code: str, name: str, data_type: str = "VARCHAR", *, required: bool = False, desc: str = "") -> dict:
    return {
        "field_code": code,
        "field_name": name,
        "data_type": data_type,
        "length": "255" if data_type == "VARCHAR" else "",
        "is_pk": False,
        "is_fk": False,
        "nullable": not required,
        "default_value": "",
        "description": desc or name,
    }


def _default_business_fields(table_name: str) -> list[dict]:
    if "巡检" in table_name or "检查" in table_name:
        return [
            _field_template("inspection_no", "巡检编号", required=True),
            _field_template("inspection_date", "巡检日期", "DATETIME", required=True),
            _field_template("inspection_location", "巡检位置", required=True),
            _field_template("inspection_item", "巡检项目", required=True),
            _field_template("inspection_result", "巡检结果", required=True),
            _field_template("abnormal_desc", "异常说明", "TEXT"),
            _field_template("handler_name", "处理人"),
            _field_template("attachment_desc", "附件说明", "TEXT"),
        ]
    if "会议" in table_name:
        return [
            _field_template("meeting_topic", "会议主题", required=True),
            _field_template("meeting_date", "会议日期", "DATETIME", required=True),
            _field_template("meeting_room", "会议室", required=True),
            _field_template("applicant_name", "报名人", required=True),
            _field_template("participant_count", "参会人数", "INT"),
            _field_template("contact_phone", "联系电话"),
            _field_template("remark", "备注", "TEXT"),
        ]
    return [
        _field_template("record_no", "业务编号", required=True),
        _field_template("title", "标题", required=True),
        _field_template("business_date", "业务日期", "DATETIME", required=True),
        _field_template("owner_name", "负责人"),
        _field_template("department_name", "所属部门"),
        _field_template("business_status", "业务状态"),
        _field_template("attachment_desc", "附件说明", "TEXT"),
        _field_template("remark", "备注", "TEXT"),
    ]


def _normalize_doc_field(field: Any) -> dict | None:
    if not isinstance(field, dict):
        return None
    code = str(field.get("field_code") or field.get("code") or "").strip()
    name = str(field.get("field_name") or field.get("name") or field.get("label") or code).strip()
    if not code or not name:
        return None
    if code.lower() in SYSTEM_FIELD_CODES or field.get("is_pk") is True:
        return None
    return {
        "field_code": code,
        "field_name": name,
        "data_type": _normalize_doc_data_type(field.get("data_type") or field.get("type"), name),
        "length": str(field.get("length") or field.get("max_length") or field.get("maxLength") or ""),
        "is_pk": False,
        "is_fk": bool(field.get("is_fk") or field.get("isFk")),
        "nullable": bool(field.get("nullable", True)),
        "default_value": str(field.get("default_value") or field.get("defaultValue") or ""),
        "description": str(field.get("description") or field.get("comment") or name),
    }


def _normalize_doc_tables(raw_tables: Any, app_info: dict[str, str]) -> list[dict]:
    tables: list[dict] = []
    for idx, raw_table in enumerate(_list_or_empty(raw_tables)):
        if not isinstance(raw_table, dict):
            continue
        table_name = str(raw_table.get("table_name") or raw_table.get("name") or "").strip()
        if not table_name:
            continue
        table_code = str(raw_table.get("table_code") or raw_table.get("code") or "").strip() or f"t_{_infer_doc_app_code(table_name)}"
        if not table_code.startswith("t_"):
            table_code = f"t_{table_code}"
        table_type = str(raw_table.get("table_type") or raw_table.get("type") or "主表").strip() or "主表"
        fields = [f for f in (_normalize_doc_field(item) for item in _list_or_empty(raw_table.get("fields"))) if f]
        seen_codes = {field["field_code"] for field in fields}
        for fallback in _default_business_fields(table_name):
            if len(fields) >= 6:
                break
            if fallback["field_code"] in seen_codes:
                continue
            fields.append(fallback)
            seen_codes.add(fallback["field_code"])
        tables.append({
            "table_code": table_code,
            "table_name": table_name,
            "table_type": table_type,
            "parent_table": str(raw_table.get("parent_table") or raw_table.get("parent_model_code") or "").strip(),
            "description": str(raw_table.get("description") or f"{table_name}的业务数据表"),
            "fields": fields,
        })

    if tables:
        return tables

    entity_name = app_info["name"].replace("系统", "").replace("平台", "").strip() or "业务"
    table_name = f"{entity_name}记录"
    table_code = f"t_{_infer_doc_app_code(table_name)}"
    return [{
        "table_code": table_code,
        "table_name": table_name,
        "table_type": "主表",
        "parent_table": "",
        "description": f"{table_name}主表",
        "fields": _default_business_fields(table_name)[:6],
    }]


def _normalize_doc_roles(raw_roles: Any) -> list[dict[str, str]]:
    roles: list[dict[str, str]] = []
    seen: set[str] = set()
    for idx, raw_role in enumerate(_list_or_empty(raw_roles)):
        if not isinstance(raw_role, dict):
            continue
        code = str(raw_role.get("role_code") or raw_role.get("code") or "").strip()
        name = str(raw_role.get("role_name") or raw_role.get("name") or code).strip()
        if not code:
            code = f"role_{idx + 1}"
        code_lower = code.lower()
        if code_lower in FORBIDDEN_ROLE_CODES:
            continue
        if any(keyword == name or keyword in name for keyword in FORBIDDEN_ROLE_NAME_KEYWORDS):
            continue
        if code_lower in seen or not name:
            continue
        seen.add(code_lower)
        roles.append({
            "role_code": code,
            "role_name": name,
            "description": str(raw_role.get("description") or f"{name}负责对应业务处理和数据维护"),
        })
    if roles:
        return roles
    return [{
        "role_code": "business_admin",
        "role_name": "业务管理员",
        "description": "维护业务配置、查看全量数据并处理异常情况",
    }]


def _normalize_doc_dicts(raw_dicts: Any, tables: list[dict]) -> list[dict]:
    dicts: list[dict] = []
    seen: set[str] = set()
    for idx, raw_dict in enumerate(_list_or_empty(raw_dicts)):
        if not isinstance(raw_dict, dict):
            continue
        code = str(raw_dict.get("dict_code") or raw_dict.get("code") or f"dict_{idx + 1}").strip()
        if code.lower() in {"approval_status", "audit_status", "process_status"} or code in seen:
            continue
        name = str(raw_dict.get("dict_name") or raw_dict.get("name") or code).strip()
        items = []
        for item_idx, raw_item in enumerate(_list_or_empty(raw_dict.get("items") or raw_dict.get("options"))):
            if isinstance(raw_item, dict):
                item_code = str(raw_item.get("item_code") or raw_item.get("code") or f"item_{item_idx + 1}").strip()
                item_name = str(raw_item.get("item_name") or raw_item.get("name") or item_code).strip()
            else:
                item_code = f"item_{item_idx + 1}"
                item_name = str(raw_item).strip()
            if item_name:
                items.append({"item_code": item_code, "item_name": item_name})
        if name and items:
            seen.add(code)
            dicts.append({"dict_code": code, "dict_name": name, "items": items})

    if dicts:
        return dicts

    result_field_name = next(
        (
            field.get("field_name", "")
            for table in tables
            for field in table.get("fields", [])
            if "结果" in field.get("field_name", "")
        ),
        "",
    )
    if result_field_name:
        return [{
            "dict_code": "dict_business_result",
            "dict_name": result_field_name,
            "items": [
                {"item_code": "normal", "item_name": "正常"},
                {"item_code": "abnormal", "item_name": "异常"},
                {"item_code": "pending", "item_name": "待处理"},
            ],
        }]
    return [{
        "dict_code": "dict_business_status",
        "dict_name": "业务状态",
        "items": [
            {"item_code": "draft", "item_name": "草稿"},
            {"item_code": "in_progress", "item_name": "处理中"},
            {"item_code": "completed", "item_name": "已完成"},
            {"item_code": "cancelled", "item_name": "已取消"},
        ],
    }]


def _primary_business_role_name(roles: list[dict[str, str]]) -> str:
    return roles[0]["role_name"] if roles else "业务管理员"


def _normalize_doc_modules(raw_modules: Any, tables: list[dict], roles: list[dict[str, str]]) -> list[dict]:
    modules: list[dict] = []
    for idx, raw_module in enumerate(_list_or_empty(raw_modules)):
        if not isinstance(raw_module, dict):
            continue
        name = str(raw_module.get("module_name") or raw_module.get("name") or "").strip()
        if not name:
            continue
        code = str(raw_module.get("module_code") or raw_module.get("code") or _infer_doc_app_code(name)).strip()
        features = []
        for feature_idx, raw_feature in enumerate(_list_or_empty(raw_module.get("features"))):
            if not isinstance(raw_feature, dict):
                continue
            feature_name = str(raw_feature.get("name") or raw_feature.get("feature_name") or f"功能点{feature_idx + 1}").strip()
            features.append({
                "name": feature_name,
                "description": str(raw_feature.get("description") or f"{feature_name}的业务操作"),
                "roles": _list_or_empty(raw_feature.get("roles")) or ["全部员工", _primary_business_role_name(roles)],
            })
        modules.append({
            "module_name": name,
            "module_code": code,
            "description": str(raw_module.get("description") or f"{name}覆盖对应业务对象的日常操作"),
            "features": features,
        })

    module_by_code = {module["module_code"]: module for module in modules}
    for table in tables:
        if str(table.get("table_type", "主表")).lower() in {"子表", "sub", "child"}:
            continue
        table_name = table["table_name"]
        code = table["table_code"].removeprefix("t_")
        if code in module_by_code:
            target = module_by_code[code]
            if len(target.get("features") or []) >= 3:
                continue
            target["features"] = (target.get("features") or []) + _default_module_features(table_name, roles)
            target["features"] = target["features"][:3]
            continue
        modules.append({
            "module_name": f"{table_name}管理",
            "module_code": code,
            "description": f"围绕{table_name}提供新增、查看、维护和处理能力。",
            "features": _default_module_features(table_name, roles),
        })
    return modules


def _default_module_features(table_name: str, roles: list[dict[str, str]]) -> list[dict]:
    admin_role = _primary_business_role_name(roles)
    return [
        {
            "name": f"新增{table_name}",
            "description": f"支持用户创建并暂存或提交{table_name}记录。",
            "roles": ["全部员工"],
        },
        {
            "name": f"查看{table_name}",
            "description": f"支持按权限范围查询、查看{table_name}详情和处理状态。",
            "roles": ["全部员工", admin_role],
        },
        {
            "name": f"维护{table_name}",
            "description": f"支持{admin_role}维护{table_name}数据并处理异常或待办事项。",
            "roles": [admin_role],
        },
    ]


def _normalize_doc_flows(raw_flows: Any, tables: list[dict], roles: list[dict[str, str]]) -> list[dict]:
    flows: list[dict] = []
    for idx, raw_flow in enumerate(_list_or_empty(raw_flows)):
        if not isinstance(raw_flow, dict):
            continue
        name = str(raw_flow.get("flow_name") or raw_flow.get("name") or "").strip()
        if not name:
            continue
        steps = []
        for step_idx, raw_step in enumerate(_list_or_empty(raw_flow.get("steps") or raw_flow.get("nodes"))):
            if not isinstance(raw_step, dict):
                continue
            steps.append({
                "step": raw_step.get("step") or raw_step.get("order") or step_idx + 1,
                "action": str(raw_step.get("action") or raw_step.get("name") or raw_step.get("label") or "").strip(),
                "role": str(raw_step.get("role") or raw_step.get("assignee") or "").strip(),
                "status": str(raw_step.get("status") or raw_step.get("result") or "").strip(),
            })
        if steps:
            flows.append({
                "flow_name": name,
                "flow_code": str(raw_flow.get("flow_code") or raw_flow.get("code") or f"flow_{idx + 1}").strip(),
                "description": str(raw_flow.get("description") or f"{name}用于串联业务提交、处理和归档。"),
                "table_code": str(raw_flow.get("table_code") or "").strip(),
                "steps": steps,
            })

    if flows:
        return flows

    existing_names = {flow["flow_name"] for flow in flows}
    for table in tables:
        if str(table.get("table_type", "主表")).lower() in {"子表", "sub", "child"}:
            continue
        flow_name = f"{table['table_name']}处理流程"
        if flow_name in existing_names:
            continue
        flows.append({
            "flow_name": flow_name,
            "flow_code": f"{table['table_code'].removeprefix('t_')}_flow",
            "description": f"{table['table_name']}从提交、处理到归档的基础业务流程。",
            "table_code": table["table_code"],
            "steps": [
                {"step": 1, "action": f"提交{table['table_name']}", "role": "全部员工", "status": "待处理"},
                {"step": 2, "action": f"处理{table['table_name']}", "role": _primary_business_role_name(roles), "status": "处理中"},
                {"step": 3, "action": "归档完成", "role": "系统", "status": "已完成"},
            ],
        })
    return flows


def _normalize_doc_permissions(raw_mappings: Any, tables: list[dict], roles: list[dict[str, str]]) -> list[dict]:
    existing_by_table: dict[str, dict] = {}
    for raw_mapping in _list_or_empty(raw_mappings):
        if not isinstance(raw_mapping, dict):
            continue
        table_code = str(raw_mapping.get("table_code") or raw_mapping.get("form_code") or "").strip()
        if not table_code:
            continue
        permissions = []
        for raw_perm in _list_or_empty(raw_mapping.get("permissions") or raw_mapping.get("rules")):
            if not isinstance(raw_perm, dict):
                continue
            role_code = str(raw_perm.get("role_code") or raw_perm.get("role") or "").strip()
            role_name = str(raw_perm.get("role_name") or raw_perm.get("name") or role_code).strip()
            if not role_code:
                continue
            operations = raw_perm.get("operations") or raw_perm.get("actions") or []
            if isinstance(operations, str):
                operations = [item.strip() for item in operations.split(",") if item.strip()]
            permissions.append({
                "role_code": role_code,
                "role_name": role_name,
                "operations": operations or ["查看"],
                "data_scope": str(raw_perm.get("data_scope") or raw_perm.get("scope") or raw_perm.get("data") or "self"),
            })
        existing_by_table[table_code] = {
            "table_code": table_code,
            "table_name": str(raw_mapping.get("table_name") or "").strip(),
            "permissions": permissions,
        }

    mappings: list[dict] = []
    for table in tables:
        if str(table.get("table_type", "主表")).lower() in {"子表", "sub", "child"}:
            continue
        table_code = table["table_code"]
        mapping = existing_by_table.get(table_code, {
            "table_code": table_code,
            "table_name": table["table_name"],
            "permissions": [],
        })
        mapping["table_name"] = mapping.get("table_name") or table["table_name"]
        permissions = mapping.get("permissions") or []
        permissions = [perm for perm in permissions if isinstance(perm, dict)]
        permissions = [perm for perm in permissions if perm.get("role_code") != "all_employee"]
        default_permissions = [{
            "role_code": "all_employee",
            "role_name": "全部员工",
            "operations": ["暂存", "新增", "查看", "编辑"],
            "data_scope": "self",
        }]
        existing_role_codes = {perm.get("role_code") for perm in permissions}
        for role in roles:
            if role["role_code"] in existing_role_codes:
                continue
            permissions.append({
                "role_code": role["role_code"],
                "role_name": role["role_name"],
                "operations": ["查看", "编辑", "导出", "查看审批历史"],
                "data_scope": "all",
            })
        mapping["permissions"] = default_permissions + permissions
        mappings.append(mapping)
    return mappings


def _derive_forms_for_doc_result(doc: dict) -> list[dict]:
    existing = doc.get("forms")
    if isinstance(existing, list) and existing:
        return existing
    try:
        from app.services.config_converter import convert_analysis_to_app_config

        return convert_analysis_to_app_config(doc).get("forms") or []
    except Exception as exc:
        logger.warning("derive forms from doc_result failed: %s", exc)
    forms: list[dict] = []
    for table in doc.get("tables", []) or []:
        if str(table.get("table_type", "主表")).lower() in {"子表", "sub", "child"}:
            continue
        table_code = table.get("table_code") or ""
        forms.append({
            "form_name": table.get("table_name") or table_code,
            "form_code": table_code,
            "model_code": table_code,
            "components": [
                {
                    "field_code": field.get("field_code"),
                    "field_name": field.get("field_name"),
                    "label": field.get("field_name"),
                    "modelField": f"{table_code}.{field.get('field_code')}",
                    "componentType": "FORM_TEXT_INPUT",
                    "required": not field.get("nullable", True),
                }
                for field in table.get("fields", []) or []
            ],
        })
    return forms


def normalize_doc_result(doc: dict, messages: list[dict[str, Any]] | None = None) -> dict:
    """把模型输出规范化为完整 SPEC，避免缺表单/流程时仍被当成完成态。"""
    source = doc if isinstance(doc, dict) else {}
    normalized = dict(source)
    app_info = _normalize_doc_app_info(source.get("app_info"), messages)
    roles = _normalize_doc_roles(source.get("roles"))
    tables = _normalize_doc_tables(source.get("tables") or source.get("models"), app_info)
    dicts = _normalize_doc_dicts(source.get("data_dictionary") or source.get("dicts"), tables)
    normalized.update({
        "app_info": app_info,
        "roles": roles,
        "data_dictionary": dicts,
        "tables": tables,
    })
    normalized["modules"] = _normalize_doc_modules(source.get("modules"), tables, roles)
    normalized["flows"] = _normalize_doc_flows(source.get("flows") or source.get("workflows"), tables, roles)
    normalized["role_table_mapping"] = _normalize_doc_permissions(
        source.get("role_table_mapping") or source.get("permissions"),
        tables,
        roles,
    )
    normalized["custom_development"] = []
    normalized["forms"] = _derive_forms_for_doc_result(normalized)
    return normalized


def is_valid_doc_result(doc: dict) -> bool:
    """完整结构校验，避免把半成品 SPEC 误存为 doc_result。"""
    if not isinstance(doc, dict):
        return False
    if not isinstance(doc.get("app_info"), dict):
        return False
    if any(not isinstance(doc.get(key), list) for key in DOC_REQUIRED_LIST_KEYS):
        return False
    main_tables = [
        table for table in doc.get("tables", [])
        if isinstance(table, dict) and str(table.get("table_type", "主表")).lower() not in {"子表", "sub", "child"}
    ]
    if not main_tables:
        return False
    if not doc.get("forms") or not doc.get("modules") or not doc.get("flows") or not doc.get("role_table_mapping"):
        return False
    return all(len(table.get("fields") or []) >= 6 for table in main_tables)


async def _save_assistant_message(session_id: int, content: str) -> None:
    """Persist an assistant message for async SSE flows that stop before doc_result."""
    if not content.strip():
        return
    from app.database import AsyncSessionLocal

    async with AsyncSessionLocal() as save_db:
        save_db.add(Message(conversation_id=session_id, role="assistant", content=content))
        await save_db.commit()


def _validation_payload(preflight) -> dict[str, Any]:
    return {
        "needs_user_input": True,
        "assistant_message": preflight.assistant_message,
        "validation": preflight.to_payload(),
    }


def _looks_like_asset_management(user_text: str) -> bool:
    keywords = ("资产", "台账", "入库", "领用", "调拨", "维修", "盘点", "报废", "折旧", "扫码")
    return "资产" in user_text and sum(1 for keyword in keywords if keyword in user_text) >= 3


def _build_asset_management_fallback_doc_result(user_text: str) -> dict:
    """资产管理场景兜底，避免模型超时时退回通用申请模板。"""
    def fields(items: list[tuple[str, str, str]]) -> list[dict]:
        return [_field_template(code, name, data_type, required=index < 2) for index, (code, name, data_type) in enumerate(items)]

    table_specs = [
        ("asset_ledger", "资产台账", "记录资产全生命周期主档和财务折旧信息", [
            ("asset_code", "资产编码", "VARCHAR"), ("asset_name", "资产名称", "VARCHAR"), ("asset_category", "资产类别", "VARCHAR"),
            ("brand_model", "品牌型号", "VARCHAR"), ("purchase_date", "购置日期", "DATE"), ("original_value", "原值", "DECIMAL"),
            ("residual_rate", "残值率", "DECIMAL"), ("depreciation_method", "折旧方式", "VARCHAR"), ("asset_status", "资产状态", "VARCHAR"),
            ("usage_status", "使用状态", "VARCHAR"), ("custodian_dept", "保管部门", "VARCHAR"), ("custodian_user", "保管人", "VARCHAR"),
            ("storage_location", "存放位置", "VARCHAR"),
        ]),
        ("asset_inbound", "资产入库单", "记录新购或批量导入资产的入库登记和确认", [
            ("inbound_no", "入库单号", "VARCHAR"), ("inbound_date", "入库日期", "DATE"), ("supplier_name", "供应商", "VARCHAR"),
            ("warehouse_name", "入库仓库", "VARCHAR"), ("asset_category", "资产类别", "VARCHAR"), ("quantity", "入库数量", "INT"),
            ("total_amount", "入库金额", "DECIMAL"), ("handler_name", "经办人", "VARCHAR"), ("confirm_status", "确认状态", "VARCHAR"),
        ]),
        ("asset_assignment", "资产领用单", "记录员工或部门领用资产的申请、审批和归还状态", [
            ("assignment_no", "领用单号", "VARCHAR"), ("asset_code", "资产编码", "VARCHAR"), ("asset_name", "资产名称", "VARCHAR"),
            ("applicant_name", "领用人", "VARCHAR"), ("applicant_dept", "领用部门", "VARCHAR"), ("apply_date", "申请日期", "DATE"),
            ("expected_return_date", "预计归还日期", "DATE"), ("assignment_reason", "领用原因", "TEXT"), ("assignment_status", "领用状态", "VARCHAR"),
        ]),
        ("asset_transfer", "资产调拨单", "记录资产跨部门或跨地点调拨及审批", [
            ("transfer_no", "调拨单号", "VARCHAR"), ("asset_code", "资产编码", "VARCHAR"), ("from_dept", "调出部门", "VARCHAR"),
            ("to_dept", "调入部门", "VARCHAR"), ("from_location", "原存放位置", "VARCHAR"), ("to_location", "新存放位置", "VARCHAR"),
            ("transfer_date", "调拨日期", "DATE"), ("transfer_reason", "调拨原因", "TEXT"), ("transfer_status", "调拨状态", "VARCHAR"),
        ]),
        ("repair_ticket", "维修工单", "记录资产报修、派单、维修处理和结果验收", [
            ("ticket_no", "工单号", "VARCHAR"), ("asset_code", "资产编码", "VARCHAR"), ("reporter_name", "报修人", "VARCHAR"),
            ("repair_type", "维修类型", "VARCHAR"), ("fault_desc", "故障描述", "TEXT"), ("report_time", "报修时间", "DATETIME"),
            ("assignee_name", "维修处理人", "VARCHAR"), ("repair_result", "维修结果", "TEXT"), ("ticket_status", "工单状态", "VARCHAR"),
        ]),
        ("inventory_task", "盘点任务", "记录盘点计划、扫码结果、差异和复核情况", [
            ("task_no", "盘点任务号", "VARCHAR"), ("task_name", "盘点任务名称", "VARCHAR"), ("inventory_scope", "盘点范围", "TEXT"),
            ("plan_start_date", "计划开始日期", "DATE"), ("plan_end_date", "计划结束日期", "DATE"), ("owner_name", "盘点负责人", "VARCHAR"),
            ("inventory_result", "盘点结果", "VARCHAR"), ("difference_desc", "差异说明", "TEXT"), ("review_status", "复核状态", "VARCHAR"),
        ]),
        ("scrap_request", "报废申请", "记录资产报废申请、部门审批、财务审核和最终处理", [
            ("scrap_no", "报废单号", "VARCHAR"), ("asset_code", "资产编码", "VARCHAR"), ("asset_name", "资产名称", "VARCHAR"),
            ("scrap_reason", "报废原因", "VARCHAR"), ("apply_date", "申请日期", "DATE"), ("original_value", "原值", "DECIMAL"),
            ("net_value", "净值", "DECIMAL"), ("applicant_dept", "申请部门", "VARCHAR"), ("scrap_status", "报废状态", "VARCHAR"),
        ]),
    ]
    tables = [
        {
            "table_code": code,
            "table_name": name,
            "table_type": "主表",
            "parent_table": "",
            "description": desc,
            "fields": fields(field_items),
        }
        for code, name, desc, field_items in table_specs
    ]
    dict_specs = [
        ("asset_category", "资产类别", ["电子设备", "办公家具", "生产设备", "交通工具", "无形资产", "其他"]),
        ("asset_status", "资产状态", ["闲置", "使用中", "调拨中", "维修中", "已报废"]),
        ("usage_status", "使用状态", ["可领用", "已领用", "已归还", "停用"]),
        ("repair_type", "维修类型", ["硬件故障", "软件维护", "定期保养", "意外损坏"]),
        ("inventory_result", "盘点结果", ["账实相符", "盘盈", "盘亏", "损毁待报"]),
        ("scrap_reason", "报废原因", ["达到使用年限", "技术淘汰", "严重损坏", "管理丢失", "其他"]),
        ("depreciation_method", "折旧方式", ["平均年限法", "工作量法", "双倍余额递减法", "年数总和法"]),
    ]
    return {
        "app_info": {
            "code": "asset_mgmt",
            "name": "企业资产管理系统",
            "description": "覆盖资产入库、台账、领用、调拨、维修、盘点、报废和折旧核算的全生命周期管理，支持扫码作业和后续财务系统同步。",
        },
        "roles": [
            {"role_code": "asset_admin", "role_name": "资产管理员", "description": "维护全公司资产台账、派单、盘点和异常处理"},
            {"role_code": "warehouse_admin", "role_name": "仓库管理员", "description": "负责资产入库、库存和仓库信息维护"},
            {"role_code": "dept_manager", "role_name": "部门负责人", "description": "审批本部门领用、调拨和报废申请，查看本部门资产"},
            {"role_code": "finance_auditor", "role_name": "财务审核员", "description": "审核报废申请并查看折旧、净值和财务同步数据"},
        ],
        "data_dictionary": [
            {
                "dict_code": code,
                "dict_name": name,
                "items": [{"item_code": re.sub(r"[^a-z0-9]+", "_", option.lower()).strip("_") or f"item_{idx + 1}", "item_name": option} for idx, option in enumerate(options)],
            }
            for code, name, options in dict_specs
        ],
        "tables": tables,
        "modules": [
            {"module_code": "asset_ledger_mgmt", "module_name": "资产台账管理", "description": "维护资产主档、状态和折旧字段", "features": [{"name": "新增资产台账", "description": "登记资产基础信息", "roles": ["资产管理员"]}, {"name": "查看资产状态", "description": "按部门、状态和类别查询资产", "roles": ["资产管理员", "部门负责人", "财务审核员"]}, {"name": "维护折旧信息", "description": "维护原值、残值率、折旧方式和净值", "roles": ["资产管理员", "财务审核员"]}]},
            {"module_code": "asset_flow_mgmt", "module_name": "资产流转管理", "description": "覆盖入库、领用、归还和调拨", "features": [{"name": "入库确认", "description": "登记并确认资产入库", "roles": ["仓库管理员", "资产管理员"]}, {"name": "领用申请", "description": "全体成员发起领用并由部门负责人审批", "roles": ["全部员工", "部门负责人"]}, {"name": "调拨审批", "description": "资产跨部门调拨审批", "roles": ["部门负责人", "资产管理员"]}]},
            {"module_code": "repair_inventory_scrap", "module_name": "维修盘点报废", "description": "覆盖维修派单、扫码盘点和报废审核", "features": [{"name": "维修派单", "description": "资产管理员派单并跟进维修结果", "roles": ["资产管理员"]}, {"name": "扫码盘点", "description": "移动端扫码完成盘点任务", "roles": ["资产管理员", "仓库管理员"]}, {"name": "报废审核", "description": "部门负责人和财务审核员分段审核", "roles": ["部门负责人", "财务审核员"]}]},
        ],
        "flows": [
            {"flow_code": "asset_inbound_confirm", "flow_name": "资产入库确认流程", "description": "仓库登记后由资产管理员确认入库并生成台账", "steps": [{"step": 1, "action": "仓库管理员提交入库单", "role": "仓库管理员", "status": "待确认"}, {"step": 2, "action": "资产管理员确认入库", "role": "资产管理员", "status": "已入库"}, {"step": 3, "action": "同步资产台账", "role": "系统", "status": "已同步"}]},
            {"flow_code": "asset_assignment_approval", "flow_name": "资产领用审批流程", "description": "全体成员发起领用，部门负责人审批后更新资产状态", "steps": [{"step": 1, "action": "员工提交领用申请", "role": "全部员工", "status": "待审批"}, {"step": 2, "action": "部门负责人审批", "role": "部门负责人", "status": "审批通过"}, {"step": 3, "action": "资产管理员确认发放", "role": "资产管理员", "status": "已领用"}]},
            {"flow_code": "repair_dispatch_flow", "flow_name": "维修工单派单流程", "description": "员工报修后资产管理员派单并确认维修结果", "steps": [{"step": 1, "action": "员工提交报修", "role": "全部员工", "status": "待派单"}, {"step": 2, "action": "资产管理员派单处理", "role": "资产管理员", "status": "维修中"}, {"step": 3, "action": "记录维修结果", "role": "维修处理人", "status": "已完成"}]},
            {"flow_code": "scrap_finance_review", "flow_name": "资产报废审核流程", "description": "报废申请先部门审批再财务审核", "steps": [{"step": 1, "action": "提交报废申请", "role": "全部员工", "status": "待部门审批"}, {"step": 2, "action": "部门负责人审批", "role": "部门负责人", "status": "待财务审核"}, {"step": 3, "action": "财务审核净值和报废原因", "role": "财务审核员", "status": "审核通过"}, {"step": 4, "action": "资产管理员执行报废", "role": "资产管理员", "status": "已报废"}]},
        ],
        "role_table_mapping": [
            {"table_code": item["table_code"], "table_name": item["table_name"], "permissions": [
                {"role_code": "all_employee", "role_name": "全部员工", "operations": ["新增", "查看"], "data_scope": "self"},
                {"role_code": "asset_admin", "role_name": "资产管理员", "operations": ["暂存", "新增", "导入", "查看", "编辑", "删除", "导出", "查看审批历史"], "data_scope": "all"},
                {"role_code": "warehouse_admin", "role_name": "仓库管理员", "operations": ["新增", "查看", "编辑", "导出"], "data_scope": "all"},
                {"role_code": "dept_manager", "role_name": "部门负责人", "operations": ["查看", "批量同意", "批量拒绝", "查看审批历史"], "data_scope": "dept"},
                {"role_code": "finance_auditor", "role_name": "财务审核员", "operations": ["查看", "批量同意", "批量拒绝", "导出"], "data_scope": "all"},
            ]} for item in tables
        ],
        "custom_development": [
            {"type": "form_component", "name": "移动端扫码盘点/扫码领用组件", "trigger": "平台标准表单无法完整覆盖摄像头扫码、资产编码解析和移动端即时回填体验", "scope": "在盘点任务、资产领用单和资产台账中实现扫码组件，回填 asset_code 并校验资产状态", "deliverables": ["扫码表单组件源码", "组件注册说明", "扫码成功/失败测试用例"], "acceptance": "移动端可扫码识别资产编码，能回填表单并提示资产不存在、已报废或不可领用等异常"},
            {"type": "hook", "name": "资产折旧计算 Hook", "trigger": "折旧涉及购置日期、原值、残值率、折旧方式和月度计算规则，超出普通字段配置", "scope": "在资产台账保存或更新时计算累计折旧、月折旧和净值，并支持后续手工重算", "deliverables": ["折旧 Hook 代码", "计算规则说明", "单元测试"], "acceptance": "输入不同折旧方式时可得到可复核的月折旧、累计折旧和净值"},
            {"type": "report", "name": "资产运营看板", "trigger": "需要跨资产台账、维修工单、盘点任务和报废申请做聚合分析", "scope": "展示资产总数、闲置率、维修中数量、即将报废资产、部门资产分布和类别占比", "deliverables": ["看板页面源码", "Mock 数据", "图表交互说明"], "acceptance": "看板可基于 mock 数据展示核心指标并支持按部门/类别筛选"},
            {"type": "integration", "name": "财务系统同步接口", "trigger": "报废和折旧数据需要推送外部财务系统，属于外部系统接口集成", "scope": "预留折旧和报废数据同步 API，支持重试、错误记录和幂等标识", "deliverables": ["同步接口模块", "API 对接说明", "失败重试测试"], "acceptance": "可用 mock 财务接口完成报废/折旧数据推送并记录同步结果"},
        ],
    }



def _build_fallback_doc_result(messages: list[dict[str, Any]]) -> dict:
    """兜底：当模型持续无法输出合法 JSON 时，返回可编辑的最小结构，保证页面可用。"""
    user_text = "\n".join((m.get("content") or "") for m in messages if m.get("role") == "user")
    if _looks_like_asset_management(user_text):
        return _build_asset_management_fallback_doc_result(user_text)

    m = re.search(r"([^\n，。；]{2,30}系统)", user_text)
    app_name = m.group(1) if m else "业务管理系统"
    app_code = "APP_" + re.sub(r"[^A-Za-z0-9]", "", app_name.upper())[:8]
    if app_code == "APP_":
        app_code = "APP_DEMO"

    table_name = app_name.replace("系统", "") + "申请"
    table_name = table_name if table_name.strip() else "业务申请"

    return {
        "app_info": {
            "code": app_code,
            "name": app_name,
            "description": "自动兜底生成的基础设计文档，请按实际需求补充和调整。"
        },
        "roles": [
            {"role_code": "applicant", "role_name": "申请人", "description": "发起业务申请"},
            {"role_code": "approver", "role_name": "审批人", "description": "审核并审批申请"},
        ],
        "data_dictionary": [],
        "tables": [
            {
                "table_code": "biz_request",
                "table_name": table_name,
                "table_type": "主表",
                "parent_table": "",
                "description": "基础业务申请主表（兜底生成）",
                "fields": [
                    {
                        "field_code": "id",
                        "field_name": "主键ID",
                        "data_type": "自增ID",
                        "length": "20",
                        "is_pk": True,
                        "is_fk": False,
                        "nullable": False,
                        "default_value": "",
                        "description": "主键",
                    },
                    {
                        "field_code": "applicant_name",
                        "field_name": "申请人",
                        "data_type": "单行输入",
                        "length": "100",
                        "is_pk": False,
                        "is_fk": False,
                        "nullable": False,
                        "default_value": "",
                        "description": "申请人姓名",
                    },
                    {
                        "field_code": "apply_date",
                        "field_name": "申请日期",
                        "data_type": "日期",
                        "length": "",
                        "is_pk": False,
                        "is_fk": False,
                        "nullable": False,
                        "default_value": "",
                        "description": "提交申请日期",
                    },
                    {
                        "field_code": "status",
                        "field_name": "状态",
                        "data_type": "单选",
                        "length": "",
                        "is_pk": False,
                        "is_fk": False,
                        "nullable": False,
                        "default_value": "draft",
                        "description": "草稿/待审批/已通过/已拒绝",
                    },
                ],
            }
        ],
        "role_table_mapping": [
            {
                "table_code": "biz_request",
                "table_name": table_name,
                "permissions": [
                    {"role_code": "all_employee", "role_name": "全部员工", "operations": ["新增", "查看"], "data_scope": "self"},
                    {"role_code": "applicant", "role_name": "申请人", "operations": ["新增", "编辑", "查看"], "data_scope": "self"},
                    {"role_code": "approver", "role_name": "审批人", "operations": ["查看", "审批"], "data_scope": "dept"},
                ],
            }
        ],
        "flows": [
            {
                "flow_code": "biz_request_flow",
                "flow_name": "业务申请审批流程",
                "description": "兜底生成的流程，可按实际业务调整",
                "table_code": "biz_request",
                "steps": [
                    {"step": 1, "action": "提交申请", "role": "申请人", "status": "待审批"},
                    {"step": 2, "action": "审批", "role": "审批人", "status": "已通过/已拒绝"},
                ],
            }
        ],
        "custom_development": [
            {
                "type": "none",
                "name": "暂无强制自开发项",
                "trigger": "兜底文档仅生成基础配置定义，当前主流程可先由低代码配置覆盖",
                "scope": "如后续补充复杂交互、外部接口或算法规则，再进入 IDE 自开发",
                "deliverables": [],
                "acceptance": "低代码配置可完成主流程演示",
            }
        ],
    }


def _normalize_custom_development_items(doc: dict) -> list[dict[str, str]]:
    source = (
        doc.get("custom_development")
        or doc.get("customDevelopment")
        or doc.get("custom_dev")
        or []
    )
    if isinstance(source, dict):
        source = source.get("items") or source.get("tasks") or source.get("features") or []
    if not isinstance(source, list):
        source = []

    items: list[dict[str, str]] = []
    for idx, item in enumerate(source):
        if not isinstance(item, dict):
            continue
        deliverables = item.get("deliverables") or item.get("deliverable") or ""
        if isinstance(deliverables, list):
            deliverables_text = "、".join(str(v) for v in deliverables if str(v).strip())
        else:
            deliverables_text = str(deliverables or "").strip()
        items.append({
            "type": str(item.get("type") or item.get("scene") or item.get("category") or "自开发扩展").strip(),
            "name": str(item.get("name") or item.get("item_name") or item.get("title") or f"自开发项 {idx + 1}").strip(),
            "trigger": str(item.get("trigger") or item.get("reason") or item.get("condition") or item.get("description") or "配置能力无法完整覆盖").strip(),
            "scope": str(item.get("scope") or item.get("implementation") or deliverables_text or "在 IDE 中实现并回写项目上下文").strip(),
            "acceptance": str(item.get("acceptance") or item.get("acceptance_criteria") or item.get("test") or "完成源码、联调和可演示验证").strip(),
        })

    return items or [{
        "type": "none",
        "name": "暂无强制自开发项",
        "trigger": "当前需求可先由模型、表单、权限和基础流程配置覆盖",
        "scope": "如后续出现复杂交互、外部接口、算法规则或报表看板，再进入 IDE 补充",
        "acceptance": "低代码配置可完成主流程演示",
    }]


def json_to_markdown(data: dict) -> str:
    """Convert requirements JSON into the standard parseable design-doc MD."""

    def cell(value: Any) -> str:
        text = str(value or "").replace("\r\n", "\n").replace("\n", "<br>")
        return text.replace("|", "\\|").strip()

    def yes_no(value: Any, default: bool = False) -> str:
        if isinstance(value, str):
            normalized = value.strip().lower()
            if normalized in {"是", "yes", "true", "1", "y"}:
                return "是"
            if normalized in {"否", "no", "false", "0", "n"}:
                return "否"
        return "是" if bool(value) else ("是" if default else "否")

    def component_label(value: Any, field: dict | None = None) -> str:
        raw = str(value or "").strip()
        mapping = {
            "FORM_TEXT_INPUT": "单行输入",
            "FORM_TEXTAREA": "多行输入",
            "FORM_SELECT": "下拉单选",
            "FORM_MULTI_SELECT": "下拉多选",
            "FORM_DATE_PICKER": "日期",
            "FORM_DATETIME_PICKER": "日期时间",
            "FORM_NUMBER_INPUT": "数字",
            "FORM_UPLOAD": "附件上传",
            "FORM_USER_SELECT": "人员选择",
            "FORM_DEPT_SELECT": "部门选择",
            "FORM_SWITCH": "开关",
            "FORM_RADIO": "单选框",
            "FORM_CHECKBOX": "复选框",
        }
        if raw in mapping:
            return mapping[raw]
        if raw:
            return raw
        data_type = str((field or {}).get("data_type") or "").upper()
        field_name = str((field or {}).get("field_name") or "")
        if data_type in {"DATE", "DATETIME"} or any(key in field_name for key in ("日期", "时间")):
            return "日期时间" if data_type == "DATETIME" else "日期"
        if data_type in {"INT", "BIGINT", "DECIMAL"}:
            return "数字"
        if data_type == "TEXT" or any(key in field_name for key in ("说明", "描述", "备注", "原因", "内容")):
            return "多行输入"
        return "单行输入"

    def scope_label(value: Any) -> str:
        return {
            "none": "无权限",
            "self": "仅本人",
            "dept": "本部门",
            "all": "全公司",
            "custom": "自定义",
            "SELF": "仅本人",
            "CURRENT_USER_DEPT": "本部门",
            "ALL": "全公司",
        }.get(str(value or "").strip(), str(value or "").strip() or "无权限")

    def permission_flag(operations: list, *keywords: str) -> str:
        ops = {str(op).strip() for op in (operations or [])}
        return "是" if any(keyword in ops for keyword in keywords) else "否"

    app_info = data.get("app_info") or {}
    tables = [table for table in (data.get("tables") or []) if isinstance(table, dict)]
    forms = [form for form in (data.get("forms") or []) if isinstance(form, dict)]
    table_by_code = {str(table.get("table_code") or "").strip(): table for table in tables}

    lines: list[str] = []
    lines.append(f"# {cell(app_info.get('name') or '功能设计文档')}")
    lines.append("")

    lines.append("## 一、应用信息")
    lines.append("")
    lines.append("| 应用名称 | 应用编码 | 说明 |")
    lines.append("|---|---|---|")
    lines.append(f"| {cell(app_info.get('name'))} | {cell(app_info.get('code'))} | {cell(app_info.get('description'))} |")
    lines.append("")

    lines.append("## 二、角色列表")
    lines.append("")
    lines.append("| 角色编码 | 角色名称 | 职责说明 |")
    lines.append("|---|---|---|")
    for role in data.get("roles") or []:
        if not isinstance(role, dict):
            continue
        lines.append(f"| {cell(role.get('role_code'))} | {cell(role.get('role_name'))} | {cell(role.get('description'))} |")
    lines.append("")

    lines.append("## 三、数据字典")
    lines.append("")
    dicts = [item for item in (data.get("data_dictionary") or []) if isinstance(item, dict)]
    if dicts:
        for dictionary in dicts:
            lines.append(f"### {cell(dictionary.get('dict_name'))}（{cell(dictionary.get('dict_code'))}）")
            lines.append("")
            lines.append("| 选项编码 | 选项名称 |")
            lines.append("|---|---|")
            for item in dictionary.get("items") or []:
                if isinstance(item, dict):
                    lines.append(f"| {cell(item.get('item_code'))} | {cell(item.get('item_name'))} |")
            lines.append("")
    else:
        lines.append("暂无数据字典。")
        lines.append("")

    lines.append("## 四、数据模型")
    lines.append("")
    lines.append("### 4.1 模型定义")
    lines.append("")
    lines.append("| 模型编码 | 模型名称 | 类型 | 所属主表模型编码 | 说明 |")
    lines.append("|---|---|---|---|---|")
    for table in tables:
        lines.append(
            "| "
            f"{cell(table.get('table_code'))} | "
            f"{cell(table.get('table_name'))} | "
            f"{cell(table.get('table_type') or '主表')} | "
            f"{cell(table.get('parent_table'))} | "
            f"{cell(table.get('description'))} |"
        )
    lines.append("")
    lines.append("### 4.2 模型字段")
    lines.append("")
    lines.append("| 模型编码 | 字段编码 | 字段名称 | 字段类型 | 数据库字段类型 | 长度/精度 | 必填 | 字典编码 | 关联模型编码 | 关联显示字段编码 | 说明 |")
    lines.append("|---|---|---|---|---|---|---|---|---|---|---|")
    for table in tables:
        table_code = str(table.get("table_code") or "").strip()
        for field in table.get("fields") or []:
            if not isinstance(field, dict):
                continue
            data_type = field.get("data_type") or field.get("database_field_type") or ""
            lines.append(
                "| "
                f"{cell(table_code)} | "
                f"{cell(field.get('field_code'))} | "
                f"{cell(field.get('field_name'))} | "
                f"{cell(component_label('', field))} | "
                f"{cell(data_type)} | "
                f"{cell(field.get('length'))} | "
                f"{yes_no(not field.get('nullable', True))} | "
                f"{cell(field.get('dict_code') or field.get('dict'))} | "
                f"{cell(field.get('ref_model') or field.get('target_model_code'))} | "
                f"{cell(field.get('ref_field') or field.get('target_field_code'))} | "
                f"{cell(field.get('description'))} |"
            )
    lines.append("")

    lines.append("## 五、表单定义")
    lines.append("")
    lines.append("### 5.1 表单清单")
    lines.append("")
    lines.append("| 表单编码 | 表单名称 | 绑定主表模型 | 说明 |")
    lines.append("|---|---|---|---|")
    if forms:
        for form in forms:
            form_code = form.get("form_code") or form.get("code") or form.get("model_code") or form.get("modelCode")
            form_name = form.get("form_name") or form.get("name") or form_code
            model_code = form.get("model_code") or form.get("modelCode") or form_code
            lines.append(f"| {cell(form_code)} | {cell(form_name)} | {cell(model_code)} | {cell(form.get('description'))} |")
    else:
        for table in tables:
            if str(table.get("table_type", "主表")).lower() in {"子表", "sub", "child"}:
                continue
            lines.append(f"| {cell(table.get('table_code'))} | {cell(table.get('table_name'))} | {cell(table.get('table_code'))} | {cell(table.get('description'))} |")
    lines.append("")

    lines.append("### 5.2 主表字段定义")
    lines.append("")
    lines.append("| 表单名称 | 字段编码 | 字段名称 | 组件类型 | 必填 | 隐藏 | 只读 | 列表展示 | 查询条件 | 字典编码 | 目标模型编码 | 目标字段编码 | 本表关联字段编码 | 说明 |")
    lines.append("|---|---|---|---|---|---|---|---|---|---|---|---|---|---|")
    if forms:
        for form in forms:
            form_name = str(form.get("form_name") or form.get("name") or form.get("form_code") or "").strip()
            model_code = str(form.get("model_code") or form.get("modelCode") or form.get("form_code") or "").strip()
            model_fields = {
                str(field.get("field_code") or "").strip(): field
                for field in (table_by_code.get(model_code, {}).get("fields") or [])
                if isinstance(field, dict)
            }
            components = form.get("components") or []
            for component in components:
                if not isinstance(component, dict):
                    continue
                section_type = str(component.get("section_type") or component.get("sectionType") or "main").lower()
                if section_type == "sub":
                    continue
                field_code = str(component.get("field_code") or component.get("code") or "").strip()
                field = model_fields.get(field_code, {})
                lines.append(
                    "| "
                    f"{cell(form_name)} | "
                    f"{cell(field_code)} | "
                    f"{cell(component.get('field_name') or component.get('label') or field.get('field_name'))} | "
                    f"{cell(component_label(component.get('component_type') or component.get('componentType'), field))} | "
                    f"{yes_no(component.get('required'))} | "
                    f"{yes_no(component.get('hidden'))} | "
                    f"{yes_no(component.get('readonly'))} | "
                    f"{yes_no(component.get('show_in_list') if 'show_in_list' in component else component.get('showInList'))} | "
                    f"{yes_no(component.get('searchable'))} | "
                    f"{cell(component.get('dict_code') or component.get('dict'))} | "
                    f"{cell(component.get('target_model_code') or component.get('targetModelCode'))} | "
                    f"{cell(component.get('target_field_code') or component.get('targetFieldCode'))} | "
                    f"{cell(component.get('origin_field_code') or component.get('originFieldCode'))} | "
                    f"{cell(component.get('description'))} |"
                )
    else:
        for table in tables:
            if str(table.get("table_type", "主表")).lower() in {"子表", "sub", "child"}:
                continue
            for field in table.get("fields") or []:
                if isinstance(field, dict):
                    lines.append(
                        f"| {cell(table.get('table_name'))} | {cell(field.get('field_code'))} | {cell(field.get('field_name'))} | "
                        f"{cell(component_label('', field))} | {yes_no(not field.get('nullable', True))} | 否 | 否 | 是 | 否 | "
                        f"{cell(field.get('dict_code') or field.get('dict'))} |  |  |  | {cell(field.get('description'))} |"
                    )
    lines.append("")

    lines.append("### 5.3 子表区域定义")
    lines.append("")
    lines.append("| 表单名称 | 子表区域名称 | 绑定模型 | 说明 |")
    lines.append("|---|---|---|---|")
    has_subtable = False
    for table in tables:
        if str(table.get("table_type", "")).lower() in {"子表", "sub", "child"}:
            has_subtable = True
            parent = str(table.get("parent_table") or "").strip()
            parent_name = next(
                (item.get("table_name") for item in tables if item.get("table_code") == parent),
                parent,
            )
            lines.append(f"| {cell(parent_name)} | {cell(table.get('table_name'))} | {cell(table.get('table_code'))} | {cell(table.get('description'))} |")
    if not has_subtable:
        lines.append("|  |  |  | 暂无子表结构 |")
    lines.append("")

    lines.append("### 5.4 子表字段定义")
    lines.append("")
    lines.append("| 表单名称 | 子表区域名称 | 字段编码 | 字段名称 | 组件类型 | 必填 | 隐藏 | 只读 | 列表展示 | 字典编码 | 说明 |")
    lines.append("|---|---|---|---|---|---|---|---|---|---|---|")
    if has_subtable:
        for table in tables:
            if str(table.get("table_type", "")).lower() not in {"子表", "sub", "child"}:
                continue
            parent = str(table.get("parent_table") or "").strip()
            parent_name = next(
                (item.get("table_name") for item in tables if item.get("table_code") == parent),
                parent,
            )
            for field in table.get("fields") or []:
                if isinstance(field, dict):
                    lines.append(
                        f"| {cell(parent_name)} | {cell(table.get('table_name'))} | {cell(field.get('field_code'))} | "
                        f"{cell(field.get('field_name'))} | {cell(component_label('', field))} | {yes_no(not field.get('nullable', True))} | "
                        f"否 | 否 | 是 | {cell(field.get('dict_code') or field.get('dict'))} | {cell(field.get('description'))} |"
                    )
    else:
        lines.append("|  |  |  |  |  | 否 | 否 | 否 | 否 |  | 暂无子表字段 |")
    lines.append("")

    lines.append("## 六、流程配置")
    lines.append("")
    flows = [flow for flow in (data.get("flows") or []) if isinstance(flow, dict)]
    if flows:
        for flow in flows:
            lines.append(f"### {cell(flow.get('flow_name'))}（{cell(flow.get('flow_code'))}）")
            lines.append("")
            if flow.get("description"):
                lines.append(cell(flow.get("description")))
                lines.append("")
            lines.append("| 步骤 | 动作 | 角色 | 状态/结果 |")
            lines.append("|---|---|---|---|")
            for step in flow.get("steps") or []:
                if isinstance(step, dict):
                    lines.append(f"| {cell(step.get('step'))} | {cell(step.get('action'))} | {cell(step.get('role'))} | {cell(step.get('status'))} |")
            lines.append("")
    else:
        lines.append("暂无流程配置。")
        lines.append("")

    lines.append("## 七、权限定义")
    lines.append("")
    lines.append("| 表单名称 | 角色编码 | 可暂存 | 可新增 | 可导入 | 可查看 | 可编辑 | 可删除 | 可导出 | 数据范围 |")
    lines.append("|---|---|---|---|---|---|---|---|---|---|")
    for mapping in data.get("role_table_mapping") or []:
        if not isinstance(mapping, dict):
            continue
        form_name = mapping.get("table_name") or mapping.get("form_name") or mapping.get("table_code")
        for permission in mapping.get("permissions") or []:
            if not isinstance(permission, dict):
                continue
            operations = permission.get("operations") or []
            lines.append(
                "| "
                f"{cell(form_name)} | "
                f"{cell(permission.get('role_code'))} | "
                f"{permission_flag(operations, '暂存')} | "
                f"{permission_flag(operations, '新增', '复制新建')} | "
                f"{permission_flag(operations, '导入')} | "
                f"{permission_flag(operations, '查看', '查看审批历史')} | "
                f"{permission_flag(operations, '编辑')} | "
                f"{permission_flag(operations, '删除', '批量删除')} | "
                f"{permission_flag(operations, '导出')} | "
                f"{scope_label(permission.get('data_scope'))} |"
            )
    lines.append("")
    lines.append("> 自开发内容不在本次配置文档中生成；后续如需复杂组件、外部接口或 Hook，请从 Vibe Coding/IDE 入口单独处理。")
    lines.append("")

    return "\n".join(lines)


def _clean_requirement_lines(text: str) -> list[str]:
    return [
        re.sub(r"^[-*•\d.\s]+", "", line).strip()
        for line in (text or "").replace("\r\n", "\n").split("\n")
        if line.strip()
    ]


def _strip_think_blocks(text: str) -> str:
    cleaned = re.sub(r"<think>[\s\S]*?</think>", "", text or "", flags=re.IGNORECASE)
    open_idx = cleaned.lower().find("<think>")
    if open_idx >= 0:
        cleaned = cleaned[:open_idx]
    return cleaned.strip()


def _join_line_summary(lines: list[str], fallback: str) -> str:
    if not lines:
        return fallback
    if len(lines) == 1:
        return lines[0]
    return "\n".join(f"- {line}" for line in lines)


def _infer_unified_app_name(business_input: str, coding_focus: str = "", uploaded_file_name: str = "") -> str:
    source = f"{business_input}\n{coding_focus}".strip()
    match = re.search(r"([^\n，。；：]{2,24}(系统|平台|应用))", source)
    if match:
        return match.group(1)
    if re.search(r"请假|休假|年假", source):
        return "请假审批系统"
    if re.search(r"报销|费用", source):
        return "费用报销系统"
    if re.search(r"采购|供应商", source):
        return "采购管理系统"
    if uploaded_file_name:
        return re.sub(r"\.[^.]+$", "", uploaded_file_name)
    return "业务管理系统"


def _infer_unified_app_code(app_name: str) -> str:
    normalized = re.sub(r"[^A-Z0-9]", "", app_name.upper())[:8]
    return f"APP_{normalized or 'DEMO'}"


def _infer_entity_name(business_input: str, coding_focus: str = "") -> str:
    source = f"{business_input}\n{coding_focus}"
    if re.search(r"请假|休假|年假", source):
        return "请假申请"
    if re.search(r"报销|费用", source):
        return "报销申请"
    if re.search(r"采购|供应商", source):
        return "采购申请"
    return "业务申请"


def _infer_recommended_scene(doc_result: dict, coding_focus: str = "") -> str:
    text_parts = [coding_focus]
    text_parts.extend(flow.get("flow_name", "") for flow in (doc_result.get("flows") or []))
    for module in (doc_result.get("modules") or []):
        for feature in (module.get("features") or []):
            text_parts.append(feature.get("name", ""))
    text_parts.extend(table.get("table_name", "") for table in (doc_result.get("tables") or []))
    raw = "\n".join(part for part in text_parts if part)

    if re.search(r"接口|API|后端|集成|同步|Webhook|定时|任务|脚本|服务", raw, re.IGNORECASE):
        return "backend"
    if re.search(r"移动|mobile|H5|小程序", raw, re.IGNORECASE) and re.search(r"组件|控件|选择器|上传|图表|卡片", raw, re.IGNORECASE):
        return "component-mobile"
    if re.search(r"移动|mobile|H5|小程序", raw, re.IGNORECASE):
        return "page-mobile"
    if re.search(r"组件|控件|上传|图表|选择器|日历|看板卡片|自定义字段组件", raw, re.IGNORECASE):
        return "component-pc"
    return "page-pc"


def _build_unified_kickoff_message(business_input: str, coding_focus: str = "", uploaded_file_name: str = "") -> str:
    blocks: list[str] = []
    if business_input.strip():
        blocks.append(f"业务需求说明：\n{business_input.strip()}")
    if coding_focus.strip():
        blocks.append(
            "补充说明：以下内容既要用于智能搭建的设计文档整理，也要作为后续智能开发的输入边界，请一起吸收：\n"
            f"{coding_focus.strip()}"
        )
    if uploaded_file_name:
        blocks.append(f"本次还上传了需求附件：{uploaded_file_name}。请先完整吸收附件内容，再整理成统一方案。")
    blocks.append("请先整理一版统一的需求摘要，后续我会继续自动生成标准设计文档，并拆出智能开发输入。")
    return "\n\n".join(blocks)


def _build_unified_fallback_summary(business_input: str, coding_focus: str = "", uploaded_file_name: str = "") -> str:
    app_name = _infer_unified_app_name(business_input, coding_focus, uploaded_file_name)
    entity_name = _infer_entity_name(business_input, coding_focus)
    return "\n".join([
        "当前租户的需求分析模型暂时不可用，已切换为本地兜底模式。",
        "",
        f"- 已根据输入整理出一版基础应用骨架：**{app_name}**",
        f"- 智能搭建侧会先围绕 **{entity_name}** 的角色、数据表、流程和权限矩阵生成标准设计文档",
        "- 智能开发侧会继续消费同一份设计文档，并结合你填写的自开发关注点生成任务简报",
        "- 建议后续在模型恢复后，再对这份基础文档做一轮精修",
    ])


def _build_unified_fallback_doc_result(business_input: str, coding_focus: str = "", uploaded_file_name: str = "") -> dict:
    app_name = _infer_unified_app_name(business_input, coding_focus, uploaded_file_name)
    entity_name = _infer_entity_name(business_input, coding_focus)
    table_code = "t_leave_request" if re.search(r"请假|休假|年假", entity_name) else "t_business_request"
    module_name = entity_name.replace("申请", "") + "管理"
    coding_lines = _clean_requirement_lines(coding_focus)

    return {
        "app_info": {
            "code": _infer_unified_app_code(app_name),
            "name": app_name,
            "description": f"{app_name}的基础兜底设计文档，建议在模型恢复后继续补充细节。",
        },
        "roles": [
            {
                "role_code": "dept_manager",
                "role_name": "部门经理",
                "description": "负责审批和查看本部门业务申请",
            },
            {
                "role_code": "hr_specialist",
                "role_name": "HR专员",
                "description": "负责查看统计、维护规则和全局台账",
            },
        ],
        "data_dictionary": [
            {
                "dict_code": "request_status",
                "dict_name": "申请状态",
                "items": [
                    {"item_code": "DRAFT", "item_name": "草稿"},
                    {"item_code": "PENDING", "item_name": "待审批"},
                    {"item_code": "APPROVED", "item_name": "已通过"},
                    {"item_code": "REJECTED", "item_name": "已拒绝"},
                ],
            }
        ],
        "tables": [
            {
                "table_code": table_code,
                "table_name": entity_name,
                "table_type": "主表",
                "parent_table": "",
                "description": f"{entity_name}主数据，记录申请主体、时间范围、原因和当前状态。",
                "fields": [
                    {
                        "field_code": "request_no",
                        "field_name": "申请单号",
                        "data_type": "VARCHAR",
                        "length": "64",
                        "is_pk": False,
                        "is_fk": False,
                        "nullable": False,
                        "default_value": "",
                        "description": "业务申请编号",
                    },
                    {
                        "field_code": "applicant_name",
                        "field_name": "申请人",
                        "data_type": "VARCHAR",
                        "length": "64",
                        "is_pk": False,
                        "is_fk": False,
                        "nullable": False,
                        "default_value": "",
                        "description": "发起申请的员工姓名",
                    },
                    {
                        "field_code": "start_date",
                        "field_name": "开始日期",
                        "data_type": "DATE",
                        "length": "",
                        "is_pk": False,
                        "is_fk": False,
                        "nullable": False,
                        "default_value": "",
                        "description": "业务开始时间",
                    },
                    {
                        "field_code": "end_date",
                        "field_name": "结束日期",
                        "data_type": "DATE",
                        "length": "",
                        "is_pk": False,
                        "is_fk": False,
                        "nullable": False,
                        "default_value": "",
                        "description": "业务结束时间",
                    },
                    {
                        "field_code": "reason",
                        "field_name": "申请原因",
                        "data_type": "TEXT",
                        "length": "",
                        "is_pk": False,
                        "is_fk": False,
                        "nullable": True,
                        "default_value": "",
                        "description": "申请原因说明",
                    },
                    {
                        "field_code": "request_status",
                        "field_name": "申请状态",
                        "data_type": "VARCHAR",
                        "length": "32",
                        "is_pk": False,
                        "is_fk": False,
                        "nullable": False,
                        "default_value": "DRAFT",
                        "description": "当前审批状态",
                    },
                ],
            }
        ],
        "role_table_mapping": [
            {
                "table_code": table_code,
                "table_name": entity_name,
                "permissions": [
                    {
                        "role_code": "all_employee",
                        "role_name": "全部员工",
                        "operations": ["暂存", "新增", "查看"],
                        "data_scope": "self",
                    },
                    {
                        "role_code": "dept_manager",
                        "role_name": "部门经理",
                        "operations": ["查看", "批量同意", "批量拒绝", "查看审批历史"],
                        "data_scope": "dept",
                    },
                    {
                        "role_code": "hr_specialist",
                        "role_name": "HR专员",
                        "operations": ["查看", "导出", "日志"],
                        "data_scope": "all",
                    },
                ],
            }
        ],
        "modules": [
            {
                "module_name": module_name,
                "module_code": "request_mgmt",
                "description": f"{entity_name}的提交、审批和查询功能",
                "features": [
                    {
                        "name": f"提交{entity_name}",
                        "description": "员工录入并提交业务申请",
                        "roles": ["全部员工"],
                    },
                    {
                        "name": f"审批{entity_name}",
                        "description": "部门经理进行审批处理",
                        "roles": ["部门经理"],
                    },
                    {
                        "name": "统计分析",
                        "description": "HR 查看全局台账和统计结果",
                        "roles": ["HR专员"],
                    },
                    *[
                        {
                            "name": f"自开发扩展 {idx + 1}",
                            "description": line,
                            "roles": ["HR专员"],
                        }
                        for idx, line in enumerate(coding_lines[:2])
                    ],
                ],
            }
        ],
        "flows": [
            {
                "flow_name": f"{entity_name}审批流程",
                "flow_code": "request_approval_flow",
                "description": f"{entity_name}从提交到审批归档的标准流程",
                "steps": [
                    {"step": 1, "action": f"员工提交{entity_name}", "role": "全部员工", "status": "待审批"},
                    {"step": 2, "action": "部门经理审批", "role": "部门经理", "status": "已通过/已拒绝"},
                    {"step": 3, "action": "HR归档并统计", "role": "HR专员", "status": "已归档"},
                ],
            }
        ],
        "custom_development": [
            {
                "type": "自开发扩展" if coding_lines else "none",
                "name": f"自开发项 {idx + 1}" if line else "暂无强制自开发项",
                "trigger": line or "当前需求可先由模型、表单、权限和基础流程配置覆盖",
                "scope": "在 IDE 中实现并回写项目上下文" if line else "如后续出现复杂交互、外部接口或算法规则，再进入 IDE 补充",
                "deliverables": ["源码文件", "联调说明", "测试用例"] if line else [],
                "acceptance": "完成源码、联调和可演示验证" if line else "低代码配置可完成主流程演示",
            }
            for idx, line in enumerate(coding_lines[:3] or [""])
        ],
    }


def _build_coding_brief(doc_result: dict, business_input: str, coding_focus: str = "", uploaded_file_name: str = "", scene_category: str = "") -> str:
    app_info = doc_result.get("app_info") or {}
    role_names = [role.get("role_name", "").strip() for role in (doc_result.get("roles") or []) if role.get("role_name")]
    table_names = [table.get("table_name", "").strip() for table in (doc_result.get("tables") or []) if table.get("table_name")]
    flow_names = [flow.get("flow_name", "").strip() for flow in (doc_result.get("flows") or []) if flow.get("flow_name")]
    features = [
        feature.get("name", "").strip()
        for module in (doc_result.get("modules") or [])
        for feature in (module.get("features") or [])
        if feature.get("name")
    ]
    coding_lines = _clean_requirement_lines(coding_focus)
    business_lines = _clean_requirement_lines(business_input)
    resolved_scene = scene_category or _infer_recommended_scene(doc_result, coding_focus)

    feature_lines = features[:8] if features else ["暂无显式功能点，请结合业务对象和流程补齐实现细节。"]
    coding_scope = coding_lines or [
        "请优先关注需要自定义页面交互、复杂组件、统计分析看板或平台外部接口集成的部分。",
        "如果需求主要是平台配置即可完成，请在实现前先确认是否真的需要自开发。",
    ]

    lines = [
        f"# {(app_info.get('name') or '业务应用')} 智能开发任务简报",
        "",
        "## 一、背景与目标",
        "",
        f"- **应用名称**：{app_info.get('name') or '未命名应用'}",
        f"- **应用编码**：{app_info.get('code') or '待补充'}",
        f"- **业务目标**：{app_info.get('description') or '待补充'}",
        f"- **推荐开发场景**：{resolved_scene}",
    ]
    if uploaded_file_name:
        lines.append(f"- **来源附件**：{uploaded_file_name}")

    lines.extend([
        "",
        "## 二、原始业务输入",
        "",
        _join_line_summary(business_lines, "暂无额外业务输入，主要以结构化设计文档为准。"),
        "",
        "## 三、与智能搭建的分工",
        "",
        "- 组织角色、数据字典、数据表、表单和基础权限矩阵由智能搭建负责。",
        "- 基础流程配置和平台内可配置规则优先走低代码能力。",
        "",
        "## 四、智能开发需要重点处理的内容",
        "",
    ])
    lines.extend(f"- {line}" for line in coding_scope)
    lines.extend([
        "",
        "## 五、业务结构参考",
        "",
        f"- **角色**：{'、'.join(role_names) if role_names else '待补充'}",
        f"- **核心业务对象**：{'、'.join(table_names) if table_names else '待补充'}",
        f"- **关键流程**：{'、'.join(flow_names) if flow_names else '待补充'}",
        "",
        "## 六、建议先实现的功能点",
        "",
    ])
    lines.extend(f"- {line}" for line in feature_lines)
    lines.extend([
        "",
        "## 七、执行要求",
        "",
        "- 以当前生成的设计文档为主，不要脱离低代码平台的数据模型和权限边界。",
        "- 如果发现某项需求纯靠平台配置即可完成，应明确指出并避免不必要的自开发。",
        "- 产出物需要能回到项目协作上下文中，便于后续多人接力和发布。",
        "",
    ])
    return "\n".join(lines)


async def parse_uploaded_file(file: UploadFile) -> str:
    """解析上传的文件内容（支持 PDF/DOCX/MD/TXT）"""
    filename = file.filename or ''
    ext = os.path.splitext(filename)[1].lower()

    content = await file.read()

    if ext in ('.md', '.txt', '.markdown'):
        return content.decode('utf-8', errors='ignore')

    # 需要临时文件处理的格式
    suffix = ext or '.tmp'
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        if ext == '.pdf':
            try:
                import pdfplumber
                parts = []
                with pdfplumber.open(tmp_path) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            parts.append(t)
                return '\n'.join(parts)
            except ImportError:
                return content.decode('utf-8', errors='ignore')

        elif ext in ('.docx', '.doc'):
            try:
                from docx import Document
                doc = Document(tmp_path)
                parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            parts.append(row_text)
                return '\n'.join(parts)
            except ImportError:
                return content.decode('utf-8', errors='ignore')
        else:
            return content.decode('utf-8', errors='ignore')
    finally:
        os.unlink(tmp_path)


# ─── Schemas ──────────────────────────────────────────────────────────────────

class ChatMessageRequest(BaseModel):
    message: str


class RequirementsSessionCreateRequest(BaseModel):
    selected_llm_config_id: Optional[int] = None


class ExportMdRequest(BaseModel):
    doc_result: dict


def _conversation_project_id(conv: Conversation) -> Optional[int]:
    return getattr(conv, "project_id", None)


# ─── Routes ──────────────────────────────────────────────────────────────────

@router.post("/sessions")
async def create_session(
    ctx: Annotated[AuthContext, Depends(get_auth_context)],
    db: Annotated[AsyncSession, Depends(get_db)],
    data: Optional[RequirementsSessionCreateRequest] = Body(default=None),
):
    """创建新的需求分析会话"""
    from app.routes.llm_configs import (
        get_active_llm_config_by_id_for_purpose,
        get_default_llm_config_id_for_purpose,
    )

    selected_llm_config_id: Optional[int] = None
    requested_model_id = data.selected_llm_config_id if data else None
    if requested_model_id is not None:
        config = await get_active_llm_config_by_id_for_purpose(
            db,
            ctx.tenant_id,
            requested_model_id,
            "builder",
        )
        if not config:
            raise HTTPException(status_code=400, detail="所选模型不可用")
        selected_llm_config_id = config.id
    else:
        selected_llm_config_id = await get_default_llm_config_id_for_purpose(
            db,
            ctx.tenant_id,
            "builder",
        )

    conv = Conversation(
        user_id=ctx.user.id,
        tenant_id=ctx.tenant_id,
        title="新需求分析",
        agent_type="requirements",
        selected_llm_config_id=selected_llm_config_id,
        status="active",
    )
    db.add(conv)
    await db.flush()

    # 添加 AI 开场白
    greeting = Message(
        conversation_id=conv.id,
        role="assistant",
        content="您好！我是您的需求分析助手，将帮助您梳理应用的业务需求。\n\n请简单描述一下您想要搭建的应用是做什么的？解决什么业务问题？"
    )
    db.add(greeting)
    await db.commit()

    return {
        "id": conv.id,
        "title": conv.title,
        "created_at": conv.created_at.isoformat(),
        "updated_at": conv.updated_at.isoformat(),
        "selected_llm_config_id": conv.selected_llm_config_id,
        "project_id": _conversation_project_id(conv),
        "doc_result": None,
        "messages": [
            {"id": greeting.id, "role": "assistant", "content": greeting.content,
             "created_at": greeting.created_at.isoformat()}
        ]
    }


@router.get("/sessions")
async def list_sessions(
    ctx: Annotated[AuthContext, Depends(get_auth_context)],
    db: Annotated[AsyncSession, Depends(get_db)]
):
    """列出当前用户的所有需求分析会话"""
    result = await db.execute(
        select(Conversation)
        .where(
            Conversation.user_id == ctx.user.id,
            Conversation.tenant_id == ctx.tenant_id,
            Conversation.agent_type == "requirements"
        )
        .order_by(Conversation.updated_at.desc())
    )
    convs = result.scalars().all()
    return [
        {
            "id": c.id,
            "title": c.title,
            "created_at": c.created_at.isoformat(),
            "updated_at": c.updated_at.isoformat(),
            "selected_llm_config_id": c.selected_llm_config_id,
            "project_id": _conversation_project_id(c),
            "has_doc": c.doc_result is not None,
        }
        for c in convs
    ]


@router.get("/sessions/{session_id}")
async def get_session(
    session_id: int,
    ctx: Annotated[AuthContext, Depends(get_auth_context)],
    db: Annotated[AsyncSession, Depends(get_db)]
):
    """获取会话详情（含全部消息和 doc_result）"""
    result = await db.execute(
        select(Conversation).where(
            Conversation.id == session_id,
            Conversation.user_id == ctx.user.id,
            Conversation.tenant_id == ctx.tenant_id,
            Conversation.agent_type == "requirements"
        )
    )
    conv = result.scalar_one_or_none()
    if not conv:
        raise HTTPException(status_code=404, detail="会话不存在")

    msgs_result = await db.execute(
        select(Message)
        .where(Message.conversation_id == session_id)
        .order_by(Message.created_at)
    )
    msgs = msgs_result.scalars().all()

    return {
        "id": conv.id,
        "title": conv.title,
        "created_at": conv.created_at.isoformat(),
        "updated_at": conv.updated_at.isoformat(),
        "selected_llm_config_id": conv.selected_llm_config_id,
        "project_id": _conversation_project_id(conv),
        "doc_result": conv.doc_result,
        "messages": [
            {"id": m.id, "role": m.role, "content": m.content,
             "created_at": m.created_at.isoformat()}
            for m in msgs
        ]
    }


@router.delete("/sessions/{session_id}")
async def delete_session(
    session_id: int,
    ctx: Annotated[AuthContext, Depends(get_auth_context)],
    db: Annotated[AsyncSession, Depends(get_db)]
):
    """删除会话"""
    result = await db.execute(
        select(Conversation).where(
            Conversation.id == session_id,
            Conversation.user_id == ctx.user.id,
            Conversation.tenant_id == ctx.tenant_id,
            Conversation.agent_type == "requirements"
        )
    )
    conv = result.scalar_one_or_none()
    if not conv:
        raise HTTPException(status_code=404, detail="会话不存在")

    # 删除消息
    msgs_result = await db.execute(
        select(Message).where(Message.conversation_id == session_id)
    )
    for msg in msgs_result.scalars().all():
        await db.delete(msg)

    await db.delete(conv)
    await db.commit()
    return {"ok": True}


@router.post("/sessions/{session_id}/chat")
async def chat(
    session_id: int,
    data: ChatMessageRequest,
    ctx: Annotated[AuthContext, Depends(get_auth_context)],
    db: Annotated[AsyncSession, Depends(get_db)]
):
    """多轮对话 — SSE 流式响应"""
    result = await db.execute(
        select(Conversation).where(
            Conversation.id == session_id,
            Conversation.user_id == ctx.user.id,
            Conversation.tenant_id == ctx.tenant_id,
            Conversation.agent_type == "requirements"
        )
    )
    conv = result.scalar_one_or_none()
    if not conv:
        raise HTTPException(status_code=404, detail="会话不存在")

    # 保存用户消息
    user_msg = Message(
        conversation_id=session_id,
        role="user",
        content=data.message
    )
    db.add(user_msg)
    await db.commit()

    # 更新会话标题（首条用户消息）
    if conv.title == "新需求分析":
        title = data.message[:30] + ("..." if len(data.message) > 30 else "")
        conv.title = title
        await db.commit()

    # 获取历史消息
    msgs_result = await db.execute(
        select(Message)
        .where(Message.conversation_id == session_id)
        .order_by(Message.created_at)
    )
    all_msgs = msgs_result.scalars().all()

    # 构建 LLM 消息
    llm_messages = [{"role": "system", "content": REQUIREMENTS_CHAT_PROMPT}]
    # 截断：最多 20000 字符
    history = [{"role": m.role, "content": m.content} for m in all_msgs]
    truncated = []
    total = 0
    for msg in reversed(history):
        length = len(msg["content"] or "")
        if total + length > 20000 and truncated:
            break
        truncated.insert(0, msg)
        total += length
    llm_messages.extend(truncated)

    # 预取租户 LLM 配置（在 db session 有效时）
    llm_cfg = await _get_conversation_llm_config(db, conv)
    logger.info("requirements chat: conv=%s, selected_llm=%s, cfg=%s",
                session_id, conv.selected_llm_config_id,
                {k: v for k, v in (llm_cfg or {}).items() if k != 'api_key'})

    async def event_generator():
        assistant_content = ""
        try:
            async for chunk in _stream_with_config(llm_cfg, llm_messages):
                chunk_data = json.loads(chunk)
                choices = chunk_data.get("choices", [])
                if choices:
                    delta = choices[0].get("delta", {})
                    text = delta.get("content", "")
                    if text:
                        assistant_content += text
                        yield {"event": "chunk", "data": json.dumps({"content": text}, ensure_ascii=False)}
        except Exception as e:
            import traceback
            logger.error("requirements chat error: %s\n%s", e, traceback.format_exc())
            yield {"event": "error", "data": json.dumps({"message": str(e) or repr(e)}, ensure_ascii=False)}
            return

        # 保存 AI 消息
        from app.database import AsyncSessionLocal
        async with AsyncSessionLocal() as save_db:
            ai_msg = Message(
                conversation_id=session_id,
                role="assistant",
                content=assistant_content
            )
            save_db.add(ai_msg)
            await save_db.commit()

        yield {"event": "done", "data": "{}"}

    return EventSourceResponse(event_generator())


@router.post("/sessions/{session_id}/chat-with-file")
async def chat_with_file(
    session_id: int,
    ctx: Annotated[AuthContext, Depends(get_auth_context)],
    db: Annotated[AsyncSession, Depends(get_db)],
    message: str = Form(default=""),
    file: Optional[UploadFile] = File(default=None)
):
    """发送含文件的消息（解析文件内容并附加到消息；图片使用 vision 多模态）"""
    _IMAGE_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.webp'}
    _MIME_MAP = {'.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                 '.gif': 'image/gif', '.webp': 'image/webp'}

    result = await db.execute(
        select(Conversation).where(
            Conversation.id == session_id,
            Conversation.user_id == ctx.user.id,
            Conversation.tenant_id == ctx.tenant_id,
            Conversation.agent_type == "requirements"
        )
    )
    conv = result.scalar_one_or_none()
    if not conv:
        raise HTTPException(status_code=404, detail="会话不存在")

    # 解析文件内容
    file_content = ""
    file_name = ""
    image_data_url = ""   # 仅图片时非空
    if file and file.filename:
        file_name = file.filename
        ext = os.path.splitext(file_name)[1].lower()
        try:
            if ext in _IMAGE_EXTS:
                raw = await file.read()
                mime = _MIME_MAP.get(ext, 'image/png')
                b64 = base64.b64encode(raw).decode()
                image_data_url = f"data:{mime};base64,{b64}"
            else:
                file_content = await parse_uploaded_file(file)
        except Exception as e:
            logger.warning("file parse error: %s", e)

    # 校验消息非空
    if not message.strip() and not file_content and not image_data_url:
        raise HTTPException(status_code=400, detail="消息不能为空")

    # 保存用户消息（含文件内容，供 generate-doc 使用；前端根据格式识别并只展示文件名）
    if file_content:
        # 文字文档：把文件内容完整存入 DB，generate-doc 时 AI 能看到
        db_content = f"[上传文件：{file_name}]\n\n{file_content}"
        if message.strip():
            db_content = f"{message.strip()}\n\n[上传文件：{file_name}]\n\n{file_content}"
    elif image_data_url:
        # 图片只存文件名（图片内容不适合存文本）
        db_content = message.strip() or f"已上传文件：{file_name}"
    else:
        db_content = message.strip() or f"已上传文件：{file_name}"
    user_msg = Message(
        conversation_id=session_id,
        role="user",
        content=db_content
    )
    db.add(user_msg)
    await db.commit()

    # 更新标题
    if conv.title == "新需求分析":
        title = (message.strip() or file_name)[:30]
        conv.title = title
        await db.commit()

    # 获取历史消息
    msgs_result = await db.execute(
        select(Message)
        .where(Message.conversation_id == session_id)
        .order_by(Message.created_at)
    )
    all_msgs = msgs_result.scalars().all()

    # 构建 LLM 消息（最后一条替换为含文件/图片的完整内容）
    # 上传文件时用专项提示词，让 AI 直接分析而非追问
    _sys_prompt = REQUIREMENTS_CHAT_PROMPT
    if file_content:
        _sys_prompt += (
            "\n\n【文件分析模式】用户已上传需求文档，文档完整内容已附在消息末尾。"
            "请直接阅读文档内容，用 2-3 句话确认已读取并简述系统类型和核心功能，"
            "不要提任何问题，不要要求用户补充信息，系统将自动基于文档内容生成功能设计文档。"
        )
    llm_messages = [{"role": "system", "content": _sys_prompt}]
    history = [{"role": m.role, "content": m.content} for m in all_msgs]

    # 构建最后一条用户消息内容
    if image_data_url:
        # 多模态内容块：文字 + 图片
        last_content: list | str = []
        if message.strip():
            last_content.append({"type": "text", "text": message.strip()})
        last_content.append({"type": "image_url", "image_url": {"url": image_data_url}})
    else:
        # 纯文本（文档内容拼接）
        last_content = message.strip()
        if file_content:
            last_content = f"[上传文件：{file_name}]\n\n{file_content}"
            if message.strip():
                last_content = f"{message.strip()}\n\n[上传文件：{file_name}]\n\n{file_content}"

    if history and history[-1]["role"] == "user":
        history[-1]["content"] = last_content

    def _content_len(c) -> int:
        if isinstance(c, str):
            return len(c)
        if isinstance(c, list):
            return sum(len(b.get("text", "")) if b.get("type") == "text" else 5000 for b in c)
        return 0

    truncated = []
    total = 0
    for msg in reversed(history):
        length = _content_len(msg["content"])
        if total + length > 30000 and truncated:
            break
        truncated.insert(0, msg)
        total += length
    llm_messages.extend(truncated)

    llm_cfg = await _get_conversation_llm_config(db, conv)

    async def event_generator():
        assistant_content = ""
        try:
            async for chunk in _stream_with_config(llm_cfg, llm_messages):
                chunk_data = json.loads(chunk)
                choices = chunk_data.get("choices", [])
                if choices:
                    delta = choices[0].get("delta", {})
                    text = delta.get("content", "")
                    if text:
                        assistant_content += text
                        yield {"event": "chunk", "data": json.dumps({"content": text}, ensure_ascii=False)}
        except Exception as e:
            import traceback
            logger.error("requirements chat-with-file error: %s\n%s", e, traceback.format_exc())
            yield {"event": "error", "data": json.dumps({"message": str(e)}, ensure_ascii=False)}
            return

        from app.database import AsyncSessionLocal
        async with AsyncSessionLocal() as save_db:
            # 去重：若最后一条已是 assistant 消息则跳过（防止 SSE 重连导致重复存储）
            last_check = await save_db.execute(
                select(Message)
                .where(Message.conversation_id == session_id)
                .order_by(Message.created_at.desc())
                .limit(1)
            )
            last_db_msg = last_check.scalar_one_or_none()
            if not last_db_msg or last_db_msg.role != "assistant":
                ai_msg = Message(
                    conversation_id=session_id,
                    role="assistant",
                    content=assistant_content
                )
                save_db.add(ai_msg)
                await save_db.commit()

        yield {"event": "done", "data": "{}"}

    return EventSourceResponse(event_generator())


@router.post("/sessions/{session_id}/generate-doc")
async def generate_doc(
    session_id: int,
    ctx: Annotated[AuthContext, Depends(get_auth_context)],
    db: Annotated[AsyncSession, Depends(get_db)]
):
    """从对话历史生成结构化功能设计文档 — SSE 流式响应"""
    result = await db.execute(
        select(Conversation).where(
            Conversation.id == session_id,
            Conversation.user_id == ctx.user.id,
            Conversation.tenant_id == ctx.tenant_id,
        )
    )
    conv = result.scalar_one_or_none()
    if not conv:
        raise HTTPException(status_code=404, detail="会话不存在")

    # 获取所有对话消息
    msgs_result = await db.execute(
        select(Message)
        .where(Message.conversation_id == session_id)
        .order_by(Message.created_at)
    )
    all_msgs = msgs_result.scalars().all()

    if not all_msgs:
        raise HTTPException(status_code=400, detail="对话内容为空，无法生成文档")

    # 构建 LLM 消息：system + 对话历史（去掉开头连续 assistant 消息）+ 生成指令
    history = [{"role": m.role, "content": m.content} for m in all_msgs]
    # 确保历史第一条是 user（跳过开头的 assistant 开场白）
    while history and history[0]["role"] == "assistant":
        history = history[1:]
    # 找到含文件内容的第一条用户消息，截断时必须保留
    file_msg = next(
        (m for m in history if m["role"] == "user" and "[上传文件：" in (m.get("content") or "")),
        None
    )
    # 截断历史，最多 25000 字符（从最新往前取，始终保留文件消息）
    truncated = []
    total = 0
    for msg in reversed(history):
        length = len(msg["content"] or "")
        if total + length > 25000 and truncated:
            break
        truncated.insert(0, msg)
        total += length
    # 若文件消息因截断被丢弃，强制插入最前面
    if file_msg and file_msg not in truncated:
        truncated.insert(0, file_msg)
    llm_messages = [
        {"role": "system", "content": "你是一位经验丰富的产品分析师，擅长从需求对话中提取结构化信息。请严格按用户指令输出 JSON，不要添加任何解释文字或 Markdown 代码块标记。"}
    ]
    llm_messages.extend(truncated)
    # 附加生成文档的指令
    llm_messages.append({"role": "user", "content": GENERATE_DOC_PROMPT})

    llm_cfg = await _get_conversation_llm_config(db, conv)
    # generate-doc 输出可能较大，但 qwen-max 最多支持 8192 输出 token，强制不超过此值
    if llm_cfg:
        llm_cfg = {**llm_cfg, "max_tokens": 8000}

    async def event_generator():
        full_text = ""
        try:
            import asyncio
            yield {"event": "progress", "data": json.dumps({"message": "正在生成结构化 SPEC"}, ensure_ascii=False)}
            full_text = await asyncio.wait_for(
                _complete_with_config(
                    llm_cfg,
                    llm_messages,
                    max_tokens=8000,
                    temperature=0.0,
                ),
                timeout=45.0,
            )
        except Exception as e:
            import traceback
            logger.error("generate-doc error: %s\n%s", e, traceback.format_exc())
            doc_result = normalize_doc_result(_build_fallback_doc_result(truncated), truncated)
            preflight = validate_design_doc_preflight(doc_result)
            if not preflight.ok:
                await _save_assistant_message(session_id, preflight.assistant_message)
                yield {"event": "validation_required", "data": json.dumps(_validation_payload(preflight), ensure_ascii=False)}
                yield {"event": "done", "data": "{}"}
                return
            from app.database import AsyncSessionLocal
            async with AsyncSessionLocal() as save_db:
                save_result = await save_db.execute(
                    select(Conversation).where(Conversation.id == session_id)
                )
                save_conv = save_result.scalar_one_or_none()
                if save_conv:
                    save_conv.doc_result = doc_result
                    await save_db.commit()
            yield {"event": "result", "data": json.dumps({"doc_result": doc_result}, ensure_ascii=False)}
            yield {"event": "done", "data": "{}"}
            return

        # 提取 JSON
        logger.info("generate-doc full_text length=%d preview=%r", len(full_text), full_text[:200])
        try:
            doc_result = extract_json(full_text)
        except Exception as e:
            logger.warning("json extract failed, trying repair: %s | full_text_len=%d", e, len(full_text))
            try:
                doc_result = await _repair_doc_json(llm_cfg, full_text)
            except Exception as repair_err:
                logger.warning(
                    "json repair failed, trying regenerate: %s | full_text_len=%d",
                    repair_err,
                    len(full_text),
                )
                try:
                    doc_result = await _regenerate_doc_json(llm_cfg, llm_messages)
                except Exception as regen_err:
                    logger.error(
                        "json regenerate failed: %s | full_text_len=%d | first200=%r",
                        regen_err,
                        len(full_text),
                        full_text[:200],
                    )
                    doc_result = _build_fallback_doc_result(truncated)
                    logger.warning("using fallback doc_result for session_id=%s", session_id)

        doc_result = normalize_doc_result(doc_result, truncated)

        if not is_valid_doc_result(doc_result):
            logger.error(
                "invalid doc_result schema | keys=%s | full_text_len=%d",
                list(doc_result.keys()) if isinstance(doc_result, dict) else type(doc_result),
                len(full_text),
            )
            doc_result = normalize_doc_result(_build_fallback_doc_result(truncated), truncated)
            logger.warning("invalid schema fallback doc_result for session_id=%s", session_id)

        preflight = validate_design_doc_preflight(doc_result)
        if not preflight.ok:
            await _save_assistant_message(session_id, preflight.assistant_message)
            yield {"event": "validation_required", "data": json.dumps(_validation_payload(preflight), ensure_ascii=False)}
            yield {"event": "done", "data": "{}"}
            return

        # 写回数据库
        from app.database import AsyncSessionLocal
        async with AsyncSessionLocal() as save_db:
            save_result = await save_db.execute(
                select(Conversation).where(Conversation.id == session_id)
            )
            save_conv = save_result.scalar_one_or_none()
            if save_conv:
                save_conv.doc_result = doc_result
                await save_db.commit()

        yield {"event": "result", "data": json.dumps({"doc_result": doc_result}, ensure_ascii=False)}
        yield {"event": "done", "data": "{}"}

    return EventSourceResponse(event_generator())


GENERATE_DOC_SUMMARY_PROMPT = """根据我们之前的对话，请帮我整理一份完整的功能设计文档摘要。

请按以下格式输出（使用 Markdown）：

## 功能设计文档

### 应用概述
简要描述应用名称、用途和核心目标。

### 角色定义
列出所有角色及其职责。

### 数据字典
列出所有枚举值/下拉选项。

### 数据模型
列出每张表的名称、用途和关键字段。

### 业务流程
描述主要的审批/业务流程。

### 权限设计
简述各角色的数据范围和操作权限。

### 后续扩展
如果存在外部接口、复杂组件、Hook 或报表看板等配置暂不覆盖的内容，只作为后续扩展说明，不进入本次配置生成。

请简洁清晰地输出，不要输出 JSON，只输出可读的 Markdown 文档。"""


@router.post("/sessions/{session_id}/generate-doc-chat")
async def generate_doc_chat(
    session_id: int,
    ctx: Annotated[AuthContext, Depends(get_auth_context)],
    db: Annotated[AsyncSession, Depends(get_db)]
):
    """
    两步流式生成设计文档:
    1. 先流式输出人类可读的设计摘要（作为对话消息）
    2. 再在后台生成结构化 JSON（作为 doc_result 事件）

    SSE 事件:
    - event: content  → 流式文字（给对话气泡显示）
    - event: doc_result → 结构化 JSON（给前端解析为卡片）
    - event: error → 错误信息
    - event: done → 完成
    """
    result = await db.execute(
        select(Conversation).where(
            Conversation.id == session_id,
            Conversation.user_id == ctx.user.id,
            Conversation.tenant_id == ctx.tenant_id,
        )
    )
    conv = result.scalar_one_or_none()
    if not conv:
        raise HTTPException(status_code=404, detail="会话不存在")

    msgs_result = await db.execute(
        select(Message)
        .where(Message.conversation_id == session_id)
        .order_by(Message.created_at)
    )
    all_msgs = msgs_result.scalars().all()
    if not all_msgs:
        raise HTTPException(status_code=400, detail="对话内容为空")

    history = [{"role": m.role, "content": m.content} for m in all_msgs]
    while history and history[0]["role"] == "assistant":
        history = history[1:]

    # Truncate history to 25000 chars
    file_msg = next(
        (m for m in history if m["role"] == "user" and "[上传文件：" in (m.get("content") or "")),
        None
    )
    truncated = []
    total = 0
    for msg in reversed(history):
        length = len(msg["content"] or "")
        if total + length > 25000 and truncated:
            break
        truncated.insert(0, msg)
        total += length
    if file_msg and file_msg not in truncated:
        truncated.insert(0, file_msg)

    llm_cfg = await _get_conversation_llm_config(db, conv)
    if llm_cfg:
        llm_cfg = {**llm_cfg, "max_tokens": 8000}

    async def event_generator():
        # ── Phase 1: Stream human-readable summary ──
        summary_messages = [
            {"role": "system", "content": "你是一位经验丰富的产品分析师。请根据对话历史整理出结构化的功能设计文档。"}
        ]
        summary_messages.extend(truncated)
        summary_messages.append({"role": "user", "content": GENERATE_DOC_SUMMARY_PROMPT})

        summary_text = ""
        try:
            async for chunk in _stream_with_config(llm_cfg, summary_messages):
                chunk_data = json.loads(chunk)
                choices = chunk_data.get("choices", [])
                if choices:
                    delta = choices[0].get("delta", {})
                    text = delta.get("content", "")
                    if text:
                        summary_text += text
                        yield {"event": "content", "data": json.dumps({"content": text}, ensure_ascii=False)}
        except Exception as e:
            logger.error("generate-doc-chat summary error: %s", e, exc_info=True)
            yield {"event": "error", "data": json.dumps({"message": str(e)}, ensure_ascii=False)}
            return

        # Save summary as assistant message
        from app.database import AsyncSessionLocal
        async with AsyncSessionLocal() as save_db:
            save_db.add(Message(
                conversation_id=session_id,
                role="assistant",
                content=summary_text,
            ))
            await save_db.commit()

        # Signal that streaming text is done
        yield {"event": "content_done", "data": "{}"}

        # ── Phase 2: Generate structured JSON ──
        json_messages = [
            {"role": "system", "content": "你是一位经验丰富的产品分析师，擅长从需求对话中提取结构化信息。请严格按用户指令输出 JSON，不要添加任何解释文字或 Markdown 代码块标记。"}
        ]
        json_messages.extend(truncated)
        json_messages.append({"role": "user", "content": GENERATE_DOC_PROMPT})

        full_text = ""
        try:
            async for chunk in _stream_with_config(llm_cfg, json_messages):
                chunk_data = json.loads(chunk)
                choices = chunk_data.get("choices", [])
                if choices:
                    delta = choices[0].get("delta", {})
                    text = delta.get("content", "")
                    if text:
                        full_text += text
        except Exception as e:
            logger.error("generate-doc-chat json error: %s", e, exc_info=True)
            yield {"event": "error", "data": json.dumps({"message": str(e)}, ensure_ascii=False)}
            return

        # Extract and validate JSON
        try:
            doc_result = extract_json(full_text)
        except Exception:
            try:
                doc_result = await _repair_doc_json(llm_cfg, full_text)
            except Exception:
                try:
                    doc_result = await _regenerate_doc_json(llm_cfg, json_messages)
                except Exception:
                    doc_result = _build_fallback_doc_result(truncated)

        doc_result = normalize_doc_result(doc_result, truncated)

        if not is_valid_doc_result(doc_result):
            doc_result = normalize_doc_result(_build_fallback_doc_result(truncated), truncated)

        preflight = validate_design_doc_preflight(doc_result)
        if not preflight.ok:
            await _save_assistant_message(session_id, preflight.assistant_message)
            yield {"event": "content", "data": json.dumps({"content": "\n\n" + preflight.assistant_message}, ensure_ascii=False)}
            yield {"event": "validation_required", "data": json.dumps(_validation_payload(preflight), ensure_ascii=False)}
            yield {"event": "done", "data": "{}"}
            return

        # Save to database
        async with AsyncSessionLocal() as save_db:
            save_result = await save_db.execute(
                select(Conversation).where(Conversation.id == session_id)
            )
            save_conv = save_result.scalar_one_or_none()
            if save_conv:
                save_conv.doc_result = doc_result
                await save_db.commit()

        yield {"event": "doc_result", "data": json.dumps({"doc_result": doc_result}, ensure_ascii=False)}
        yield {"event": "done", "data": "{}"}

    return EventSourceResponse(event_generator())


@router.post("/unified-plan")
async def generate_unified_plan(
    ctx: Annotated[AuthContext, Depends(get_auth_context)],
    db: Annotated[AsyncSession, Depends(get_db)],
    business_input: str = Form(default=""),
    coding_focus: str = Form(default=""),
    selected_llm_config_id: Optional[int] = Form(default=None),
    project_id: Optional[int] = Form(default=None),
    file: Optional[UploadFile] = File(default=None),
):
    if not business_input.strip() and not (file and file.filename):
        raise HTTPException(status_code=400, detail="请先输入业务需求或上传需求文档")

    file_name = file.filename if file and file.filename else ""
    file_content = ""
    if file_name:
        try:
            ext = os.path.splitext(file_name)[1].lower()
            if ext not in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
                file_content = await parse_uploaded_file(file)
        except Exception as exc:
            logger.warning("unified-plan parse file failed: %s", exc)

    from app.routes.llm_configs import (
        get_active_llm_config_by_id_for_purpose,
        get_default_llm_config_id_for_purpose,
    )

    resolved_model_id: Optional[int] = None
    if selected_llm_config_id is not None:
        config = await get_active_llm_config_by_id_for_purpose(
            db,
            ctx.tenant_id,
            selected_llm_config_id,
            "builder",
        )
        if not config:
            raise HTTPException(status_code=400, detail="所选模型不可用")
        resolved_model_id = config.id
    else:
        resolved_model_id = await get_default_llm_config_id_for_purpose(
            db,
            ctx.tenant_id,
            "builder",
        )

    kickoff_message = _build_unified_kickoff_message(
        business_input=business_input,
        coding_focus=coding_focus,
        uploaded_file_name=file_name,
    )
    stored_user_content = kickoff_message
    if file_content:
        stored_user_content = f"{stored_user_content}\n\n[上传文件：{file_name}]\n\n{file_content}"
    elif file_name:
        stored_user_content = f"{stored_user_content}\n\n[上传文件：{file_name}]"

    title_seed = business_input.strip() or file_name or "统一 AI Builder"
    conv = Conversation(
        user_id=ctx.user.id,
        tenant_id=ctx.tenant_id,
        title=title_seed[:30] + ("..." if len(title_seed) > 30 else ""),
        agent_type="requirements",
        project_id=project_id,
        selected_llm_config_id=resolved_model_id,
        status="active",
    )
    db.add(conv)
    await db.flush()

    db.add(Message(
        conversation_id=conv.id,
        role="user",
        content=stored_user_content,
    ))
    await db.commit()

    history = [{"role": "user", "content": stored_user_content}]
    llm_cfg = await _get_conversation_llm_config(db, conv)
    summary_text = ""
    doc_result: dict | None = None
    used_fallback = False
    fallback_reason: str | None = None

    try:
        summary_messages = [
            {"role": "system", "content": "你是一位经验丰富的产品分析师。请根据对话历史整理出结构化的功能设计文档。"},
            *history,
            {"role": "user", "content": GENERATE_DOC_SUMMARY_PROMPT},
        ]
        summary_text = _strip_think_blocks(
            await _complete_with_config(llm_cfg, summary_messages, max_tokens=4000, temperature=0.2)
        )

        json_messages = [
            {
                "role": "system",
                "content": "你是一位经验丰富的产品分析师，擅长从需求对话中提取结构化信息。请严格按用户指令输出 JSON，不要添加任何解释文字或 Markdown 代码块标记。",
            },
            *history,
            {"role": "user", "content": GENERATE_DOC_PROMPT},
        ]

        full_text = await _complete_with_config(llm_cfg, json_messages, max_tokens=8000, temperature=0.0)
        try:
            doc_result = extract_json(full_text)
        except Exception:
            try:
                doc_result = await _repair_doc_json(llm_cfg, full_text)
            except Exception:
                doc_result = await _regenerate_doc_json(llm_cfg, json_messages)

        doc_result = normalize_doc_result(doc_result, history)

        if not is_valid_doc_result(doc_result):
            raise ValueError("生成的设计文档结构不完整")

    except Exception as exc:
        used_fallback = True
        fallback_reason = str(exc)
        logger.warning("unified-plan fallback for tenant=%s conv=%s: %s", ctx.tenant_id, conv.id, exc)
        summary_text = _build_unified_fallback_summary(
            business_input=business_input,
            coding_focus=coding_focus,
            uploaded_file_name=file_name,
        )
        doc_result = _build_unified_fallback_doc_result(
            business_input=business_input,
            coding_focus=coding_focus,
            uploaded_file_name=file_name,
        )

    if not summary_text.strip():
        summary_text = _build_unified_fallback_summary(
            business_input=business_input,
            coding_focus=coding_focus,
            uploaded_file_name=file_name,
        )

    if not doc_result:
        used_fallback = True
        if not fallback_reason:
            fallback_reason = "未能生成结构化设计文档"
        doc_result = _build_unified_fallback_doc_result(
            business_input=business_input,
            coding_focus=coding_focus,
            uploaded_file_name=file_name,
        )

    doc_result = normalize_doc_result(doc_result, history)

    preflight = validate_design_doc_preflight(doc_result)
    if not preflight.ok:
        db.add(Message(
            conversation_id=conv.id,
            role="assistant",
            content=preflight.assistant_message,
        ))
        await db.commit()
        return {
            "session_id": conv.id,
            "project_id": conv.project_id,
            "summary": preflight.assistant_message,
            "doc_result": None,
            "builder_markdown": "",
            "coding_brief": "",
            "recommended_scene": None,
            "used_fallback": used_fallback,
            "fallback_reason": fallback_reason,
            "needs_user_input": True,
            "validation": preflight.to_payload(),
            "source_file_name": file_name or None,
        }

    recommended_scene = _infer_recommended_scene(doc_result, coding_focus)
    builder_markdown = json_to_markdown(doc_result)
    coding_brief = _build_coding_brief(
        doc_result=doc_result,
        business_input=business_input,
        coding_focus=coding_focus,
        uploaded_file_name=file_name,
        scene_category=recommended_scene,
    )

    save_result = await db.execute(
        select(Conversation).where(Conversation.id == conv.id)
    )
    save_conv = save_result.scalar_one_or_none()
    if save_conv:
        save_conv.doc_result = doc_result
        save_conv.title = (doc_result.get("app_info") or {}).get("name") or save_conv.title

    db.add(Message(
        conversation_id=conv.id,
        role="assistant",
        content=summary_text,
    ))
    await db.commit()

    return {
        "session_id": conv.id,
        "project_id": conv.project_id,
        "summary": summary_text,
        "doc_result": doc_result,
        "builder_markdown": builder_markdown,
        "coding_brief": coding_brief,
        "recommended_scene": recommended_scene,
        "used_fallback": used_fallback,
        "fallback_reason": fallback_reason,
        "source_file_name": file_name or None,
    }


@router.post("/export-md")
async def export_md(
    data: ExportMdRequest,
    ctx: Annotated[AuthContext, Depends(get_auth_context)]
):
    """将 5 模块 JSON 转换为 Markdown 字符串"""
    try:
        md = json_to_markdown(data.doc_result)
        return {"markdown": md}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Markdown 生成失败: {e}")
