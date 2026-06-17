from docx import Document

TITLE = "示例文档"
SECTIONS = [("第一节", "这里是正文段落。"), ("第二节", "更多内容。")]

doc = Document()
doc.add_heading(TITLE, level=0)
for h, body in SECTIONS:
    doc.add_heading(h, level=1)
    doc.add_paragraph(body)
doc.save("output.docx")
print("saved output.docx")
